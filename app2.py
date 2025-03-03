import streamlit as st
import pandas as pd
import zipfile
from fpdf import FPDF
from docx import Document
from lxml import etree
import io
import re
import itertools
import numpy as np

# Função para extrair conteúdo dos tags no documento
def extract_content_controls(docx_path, target_tags):
    with zipfile.ZipFile(docx_path, 'r') as docx_zip:
        xml_content = docx_zip.read('word/document.xml')
    root = etree.XML(xml_content)
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    extracted_data = {tag: [] for tag in target_tags}
    for sdt in root.xpath('.//w:sdt', namespaces=ns):
        tag_element = sdt.find('.//w:tag', namespaces=ns)
        text_element = sdt.find('.//w:t', namespaces=ns)
        if tag_element is not None and text_element is not None:
            tag_value = tag_element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
            if tag_value in target_tags:
                extracted_data[tag_value].append(text_element.text)
    return extracted_data

# Função para extrair tabelas específicas
def extract_specific_table(docx_file, search_text):
    doc = Document(docx_file)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if search_text.lower() in cell.text.strip().lower():  
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        table_data.append(row_data)
                    return table_data 
    return None        


# Streamlit interface
st.title("Bench Plan Generator for Method Validation")

# File upload
docx_file = st.file_uploader("Upload Word Document", type=["docx"])

if docx_file:
    # Save the uploaded file
    file_path = f"/tmp/{docx_file.name}"
    with open(file_path, "wb") as f:
        f.write(docx_file.getbuffer())
    
    # Extracting 'std' data
    target_tags = ["std", "std_type"]
    std_data = extract_content_controls(docx_file, target_tags)
    df_std = pd.DataFrame({"stds": std_data["std"],
                        "std_type": std_data["std_type"]})

    # Extracting 'sam' data
    target_tags = ["sam"]
    sam_data = extract_content_controls(docx_file, target_tags)
    df_sam = pd.DataFrame(sam_data['sam'], columns=["samples"])
    
    # Extracting 'excip' data
    target_tags = ["excip"]
    excip_data = extract_content_controls(docx_file, target_tags)
    df_excip = pd.DataFrame(excip_data['excip'], columns=["sample"])
    
    # Extracting composition data
    target_tags = ["composition_name", "composition", "composition_per"]
    composition_data = extract_content_controls(docx_file, target_tags)
    min_length = min(len(composition_data["composition_name"]), len(composition_data["composition"]), len(composition_data["composition_per"]))
    df_composition = pd.DataFrame({
        "sample": composition_data["composition_name"][:min_length],
        "composition": composition_data["composition"][:min_length],
        "ratio": composition_data["composition_per"][:min_length]
    })
    target_tags = ["TC_std_name", "TC_std_level", "TC_std_relation", "TC_std_unit"]
    TC_std_data = extract_content_controls(docx_file, target_tags)
    TC_std_df = pd.DataFrame({
        "TC_std_name": TC_std_data["TC_std_name"],
        "TC_std_level": TC_std_data["TC_std_level"],
        "TC_std_relation": TC_std_data["TC_std_relation"],
        "TC_std_unit": TC_std_data["TC_std_unit"]
    })
    search_text = "in relation to"
    TC_std_df_w = extract_specific_table(docx_file, search_text)

    if TC_std_df_w:
        for row in TC_std_df_w:
            print(row)
    else:
        print("Tabela não encontrada.")
    df = pd.DataFrame(TC_std_df_w[1:], columns=TC_std_df_w[0])
    df.columns = ['Name', 'TC_std', 'Concentration']
    TC_std_df_w = df[['TC_std']]
    TC_std_df = pd.concat([TC_std_df, TC_std_df_w], axis=1)
    def to_numeric(value):
        try:
            return pd.to_numeric(value, errors='coerce')  # 'coerce' converte valores inválidos para NaN
        except ValueError:
            return None
    TC_std_df['TC_std'] = pd.to_numeric(TC_std_df['TC_std'], errors='coerce')
    target_tags = ["imp_spec_unit", "imp_loq_unit", "imp_lod_unit"]
    imp_data = extract_content_controls(docx_file, target_tags)
    imp_df = pd.DataFrame({
        "imp_spec_unit": imp_data["imp_spec_unit"],
        "imp_loq_unit": imp_data["imp_loq_unit"],
        "imp_lod_unit": imp_data["imp_lod_unit"]
    })
    search_text = "Impurities Specification/limit"  # Substitua pelo texto que você quer buscar
    imp_df_w = extract_specific_table(docx_file, search_text)
    if imp_df_w:
        for row in imp_df_w:
            print(row)
    else:
        print("Tabela não encontrada.")
    df = pd.DataFrame(imp_df_w[1:], columns=imp_df_w[0])
    df.columns = ['imp_name', 'imp_spec', 'imp_spec_unit', 'imp_loq', 'imp_loq_unit', 'imp_lod', 'imp_lod_unit']
    imp_df_w = df[['imp_name','imp_spec', 'imp_loq', 'imp_lod']]
    imp_df = pd.concat([imp_df_w, imp_df], axis=1)
    imp_df['imp_spec'] = pd.to_numeric(imp_df['imp_spec'], errors='coerce')
    imp_df['imp_loq'] = pd.to_numeric(imp_df['imp_loq'], errors='coerce')
    imp_df['imp_lod'] = pd.to_numeric(imp_df['imp_lod'], errors='coerce')
    imp_df.info()
    target_tags = ["imp_spec_u_unit"]

    imp_df_u_data = extract_content_controls(docx_file, target_tags)
    imp_df_u = pd.DataFrame({
        "imp_spec_u_unit": imp_df_u_data["imp_spec_u_unit"]
    })
    search_text = "Any unknown related substances" 

    imp_df_u_w = extract_specific_table(docx_file, search_text)
    df = pd.DataFrame(imp_df_u_w[1:], columns=imp_df_u_w[0])
    df.columns = ['Name', 'imp_spec', 'imp_spec_unit']
    imp_df_u_w = df[['Name', 'imp_spec']]
    imp_df_u = pd.concat([imp_df_u, imp_df_u_w], axis=1)
    imp_df_u['imp_spec'] = pd.to_numeric(imp_df_u['imp_spec'], errors='coerce')
    imp_df_u.info()
    target_tags = ["sam_tc_unit", "sam_tc_app"]

    sam_tc_df_data = extract_content_controls(docx_file, target_tags)

    sam_tc_df = pd.DataFrame({
        "sam_tc_unit": sam_tc_df_data["sam_tc_unit"],
        "sam_tc_app": sam_tc_df_data["sam_tc_app"]
    })
    search_text = "Sample Concentration" 
    sam_tc_df_w = extract_specific_table(docx_file, search_text)
    df = pd.DataFrame(sam_tc_df_w[1:], columns=sam_tc_df_w[0])

    df.columns = ['Name', 'sam_tc', 'sam_tc_unit', 'sam_tc_app']

    sam_tc_df_w = df[['Name', 'sam_tc']]
    sam_tc_df = pd.concat([sam_tc_df_w, sam_tc_df], axis=1)
    sam_tc_df['sam_tc'] = pd.to_numeric(sam_tc_df['sam_tc'], errors='coerce')
    
    target_tags = ["excip_tc_unit"]

    excip_tc_df_data = extract_content_controls(docx_file, target_tags)

    excip_tc_df = pd.DataFrame({
        "excip_tc_unit": excip_tc_df_data["excip_tc_unit"]
    })
    search_text = "Excipients/Placebo Concentration" 
    excip_tc_df_w = extract_specific_table(docx_file, search_text)
    df = pd.DataFrame(excip_tc_df_w[1:], columns=excip_tc_df_w[0])

    df.columns = ['Name', 'excip_tc', 'excip_tc_unit', 'df_sam']

    excip_tc_df_w = df[['Name', 'excip_tc','df_sam']]
    excip_tc_df = pd.concat([excip_tc_df_w, excip_tc_df], axis=1)
    excip_tc_df['excip_tc'] = pd.to_numeric(excip_tc_df['excip_tc'], errors='coerce')
    
    
    # Extraction for 'Selectivity' - this is the part where the file content is processed
    def extraction_parameter(docx_file, word):
        """
        Function to find a specific word in a cell and divide the content in paragraphs.
        
        :param docx_file: document file (uploaded as BytesIO object)
        :param word: word to be found
        
        :return: DataFrame with extracted paragraphs.
        """
        if not docx_file or not isinstance(docx_file, io.BytesIO):
            return None  # Ensure the file is valid

        try:
            # Use io.BytesIO to load file content
            doc = Document(io.BytesIO(docx_file.getvalue()))  
        except Exception as e:
            print(f"Error processing DOCX: {e}")
            return None

        dados = []
        for tabela in doc.tables:
            for linha in tabela.rows:
                for i, celula in enumerate(linha.cells):
                    texto_celula = celula.text.strip()
                    if word in texto_celula:
                        if i + 1 < len(linha.cells):
                            texto_direita = linha.cells[i + 1].text.strip()
                            paragrafos = texto_direita.split("\n")
                            for par in paragrafos:
                                dados.append([par.strip()])
        df = pd.DataFrame(dados, columns=["Name_sol"])
        df = df.dropna()
        return df

    # Extract 'Selectivity' Data
    df_selectivity = extraction_parameter(docx_file, 'Selectivity')
    if df_selectivity is not None and not df_selectivity.empty:
        st.write("Selectivity data extracted")
    else:
        st.warning("Error processing the file or 'Selectivity' not found.")
    
    # Drop the first index row and reset the index
    if df_selectivity is not None and not df_selectivity.empty:
        # Strip leading/trailing spaces from the 'Name_sol' column and filter out rows that are empty
        df_selectivity = df_selectivity[df_selectivity["Name_sol"].str.strip() != ""]
        # Drop the first row (index 0) and reset the index
        df_sel = df_selectivity.drop(index=0).reset_index(drop=True)
        st.write("Selectivity data processed")
    else:
        st.warning("Selectivity data could not be extracted or is empty.")

    # Extracting selectivity content into structured columns
    def extract_sel(df):
        """
        Function to divide selectivity content into columns: Name, tc, app, spiked, and tc_spiking.
        """
        def slipt_rows(linha):
            name_match = re.match(r"^(.*?)\sat", linha)
            tc_match = re.search(r"(at\s(?:method|sample|specification)\s.*\sconcentration)", linha)
            app_match = re.search(r"for\s(.*)", linha)
            spiked_match = re.search(r"spiked\s(.*?)\sat", linha)
            
            # 'tc_spiking' - The remaining part of the sentence after the last extraction
            tc_spiking_match = re.search(r"at\s.*\s(.*)", linha)

            # Extraction of the parts, if they exist
            name = name_match.group(1) if name_match else linha
            tc = tc_match.group(1) if tc_match else ""
            app = app_match.group(1) if app_match else ""
            spiked = spiked_match.group(1) if spiked_match else ""
            
            # Verificação adicional para evitar erro caso tc_spiking_match seja None
            tc_spiking = tc_spiking_match.group(1) if tc_spiking_match and tc_spiking_match.group(1) else ""

            return name, tc, app, spiked, tc_spiking

        # Apply the function to each row in the DataFrame
        df[['Name', 'tc', 'app', 'spiked', 'tc_spiking']] = df['Name_sol'].apply(lambda x: pd.Series(slipt_rows(x)))
        return df

    df_sel = extract_sel(df_sel)
    df_sel.loc[~df_sel['Name_sol'].str.contains('spiked', case=False), 'tc_spiking'] = None
    def assign_column_value(row):
        if row['Name'] == 'Standard':
            return 'std'
        elif row['Name'] in imp_df['imp_name'].values:
            return 'imp'
        elif row['Name'] in df_sam['samples'].values:
            return 'sam'
        elif row['Name'] in df_excip['sample'].values:
            return 'excip'
        elif row['Name'] == 'Resolution':
            return 'RES'
        elif 'blank' in row['Name'].lower():  # Check if 'blank' is in the Name (case insensitive)
            return 'BLK'
        else:
            return None
    
    df_sel['Category'] = df_sel.apply(assign_column_value, axis=1)
    def assign_column_value_app(row):
        if row['app'] == 'Standard':
            return 'std'
        elif row['app'] in imp_df['imp_name'].values:
            return 'imp'
        elif row['app'] in df_sam['samples'].values:
            return 'sam'
        elif row['app'] in df_excip['sample'].values:
            return 'excip'
        elif row['app'] == 'Resolution':
            return 'RES'
        elif row['app'] == '':
            return 'BLK'
        else:
            return None 
    df_sel['Category_app'] = df_sel.apply(assign_column_value_app, axis=1)
    def update_tc(row):
        if row['Category'] == 'std':
            tc_value = TC_std_df[~TC_std_df['TC_std_level'].isin(['LOQ', 'LOD'])]['TC_std']
            if not tc_value.empty:
                return tc_value.iloc[0]
            else:
                return None  

        elif row['Category'] == 'imp':
            tc_value = imp_df[imp_df['imp_name'] == row['Name']]['imp_spec']
            if not tc_value.empty:
                return tc_value.iloc[0]
            else:
                return None

        elif row['Category'] == 'excip':
            tc_value = excip_tc_df[(excip_tc_df['Name'] == row['Name']) & 
                                    (excip_tc_df['df_sam'] == row['app'])]['excip_tc']
            if not tc_value.empty:
                return tc_value.iloc[0]
            else:
                return None

        elif row['Category'] == 'sam':
            tc_value = sam_tc_df[sam_tc_df['Name'] == row['Name']]['sam_tc']
            if not tc_value.empty:
                return tc_value.iloc[0]
            else:
                return None

        return None
    df_sel['tc'] = df_sel.apply(update_tc, axis=1)
    def modify_spiked(spiked_value):
        modified_value = spiked_value.replace('with ', '').replace('and ', ', ')
        
        return [item.strip() for item in modified_value.split(',')]

    df_sel['spiked'] = df_sel['spiked'].apply(modify_spiked)
    def modify_tc_spiking(row):
        if row['tc_spiking'] == 'specification':
            values = []
            for item in row['spiked']:
                imp_spec_value = imp_df.loc[imp_df['imp_name'] == item, 'imp_spec']
                if not imp_spec_value.empty:
                    values.append(float(imp_spec_value.iloc[0]))
                else:
                    values.append(None)
            
            return values 

        elif isinstance(row['tc_spiking'], str) and '%' in row['tc_spiking']:
            try:
                value = float(row['tc_spiking'].replace('%', ''))
                return [value] * len(row['spiked'])
            except ValueError:
                return [None] * len(row['spiked'])

        elif pd.notnull(row['tc_spiking']) and row['tc_spiking'] != '':
            try:
                return [float(row['tc_spiking'])] * len(row['spiked'])
            except ValueError:
                return [None] * len(row['spiked'])
        return [None] * len(row['spiked'])
    df_sel['tc_spiking'] = df_sel.apply(modify_tc_spiking, axis=1)
    def extraction_accuracy(doc_path):
        """
        Function to locate Accuracy and extract paragraohs tant contains "For known" and "For Unknwon" withou include them.
        :param doc_path: document path.
        :return: DataFrame with paragraohs.
        """
        if not docx_file or not isinstance(docx_file, io.BytesIO):
            return None  # Ensure the file is valid

        try:
            # Use io.BytesIO to load file content
            doc = Document(io.BytesIO(docx_file.getvalue()))  
        except Exception as e:
            print(f"Error processing DOCX: {e}")
            return None
        dados = []
        extracting = False

        for tabela in doc.tables:
            for linha in tabela.rows:
                for i, celula in enumerate(linha.cells):
                    texto_celula = celula.text.strip()
                    if "Accuracy" in texto_celula:
                        print(f"Encontrado 'Accuracy' na célula: {texto_celula}")
                        if i + 1 < len(linha.cells):  
                            texto_direita = linha.cells[i + 1].text.strip()
                            print(f"Conteúdo da célula à direita: {texto_direita}")
                            paragrafos = texto_direita.split("\n")
                            for par in paragrafos:
                                par = par.strip()
                                if "for known" in par.lower(): 
                                    extracting = True
                                    print(f"Iniciando extração ao encontrar: {par}")
                                if "for unknown" in par.lower():
                                    extracting = False
                                    print(f"Interrompendo extração ao encontrar: {par}")
                                    break
                                if extracting:
                                    print(f"Adicionando parágrafo: {par}")
                                    dados.append([par])
                        else:
                            print("Não há célula à direita.")
        if not dados:
            print("Nenhum dado foi extraído.")
        else:
            print(f"{len(dados)} parágrafos extraídos.")

        df = pd.DataFrame(dados, columns=["Name_sol"])
        df = df.dropna()
        
        return df
      
    # Extract 'Selectivity' Data
    df_accuracy = extraction_accuracy(docx_file)
    if df_accuracy is  not None and not df_accuracy.empty:
        st.write("Extracted Accuracy Data")  # Display DataFrame if it contains data
    else:
        st.warning("Error processing the file or 'Selectivity' not found.")
    df_accuracy = df_accuracy[df_accuracy["Name_sol"].str.strip() != ""]
    # Drop the first index row and reset
    if df_accuracy is not None and not df_accuracy.empty:
        df_accuracy = df_accuracy.drop(index=0).reset_index(drop=True)
        st.write("Processed Accuracy Data")  # Show processed selectivity data
    else:
        st.warning("Acccuracy data could not be processed.")
    #df_accuracy = pd.DataFrame(df_accuracy)
    def process_paragraphs(df):
        df['Level'] = None
        df['sam'] = None
        df['acc_n'] = None
        df['acc_imp'] = None
        df['acc_imp_level'] = None

        for index, row in df.iterrows():
            texto = row['Name_sol']
            
            if "Level:" in texto:
                df.at[index, 'Level'] = texto.split(":")[1].strip()
            
            # Preenchendo a coluna 'sam' com base em "For" até "samples (s),"
            if texto.startswith("For") and "samples (s)," in texto:
                sam_match = re.search(r"^For\s+(.*?)\s+samples \(s\),", texto)
                if sam_match:
                    df.at[index, 'sam'] = sam_match.group(1)
            
            # Preenchendo a coluna 'acc_imp' com base em "spiked with" até "at"
            if "spiked with" in texto and "at" in texto:
                acc_imp_match = re.search(r"spiked with(.*?)at", texto)
                if acc_imp_match:
                    df.at[index, 'acc_imp'] = acc_imp_match.group(1).strip()
            
            # Preenchendo a coluna 'acc_imp_level' com base em "at" até o final
            if "at" in texto:
                acc_imp_level_match = re.search(r"at(.*)", texto)
                if acc_imp_level_match:
                    df.at[index, 'acc_imp_level'] = acc_imp_level_match.group(1).strip()
            
            # Preenchendo a coluna 'acc_n' com base nas novas condições
            if "Prepare" in texto and "as is" in texto:
                # Quando for "as is", extrai de "Prepare" até o primeiro parêntese ")"
                acc_n_match = re.search(r"Prepare(.*?[\)])", texto)
                if acc_n_match:
                    df.at[index, 'acc_n'] = acc_n_match.group(1).strip()
            elif "prepare" in texto:
                # Quando for "Prepare", extrai de "Prepare" até "independent", sem incluir "prepared" e "independent"
                acc_n_match = re.search(r"prepare(.*?)independent", texto, re.IGNORECASE)
                if acc_n_match:
                    df.at[index, 'acc_n'] = acc_n_match.group(1).strip()
            # Preenchendo a coluna 'sam' com base nas novas condições
            if "as is" in texto:
                # Extrai o conteúdo entre o primeiro parêntese ")" e "sample", sem incluir "sample"
                sam_match = re.search(r"\)(.*?)sample", texto)
                if sam_match:
                    df.at[index, 'sam'] = sam_match.group(1).strip()
            # Caso o texto não se encaixe em nenhuma das condições,  adicionar nas colunas restantes
            if pd.isna(df.at[index, 'Level']) and pd.isna(df.at[index, 'sam']) and pd.isna(df.at[index, 'acc_n']) and pd.isna(df.at[index, 'acc_imp']) and pd.isna(df.at[index, 'acc_imp_level']):
                df.at[index, 'sam'] = texto
            # Para a coluna acc_n, extrair apenas o número dentro dos parênteses
            if df.at[index, 'acc_n']:
                acc_n_number = re.search(r"\((\d+)\)", df.at[index, 'acc_n'])
                if acc_n_number:
                    df.at[index, 'acc_n'] = acc_n_number.group(1) 
            # Para a coluna 'sam' e 'acc_imp', substituir "and" por vírgula e transformar em lista
            if df.at[index, 'sam']:
                # Substitui "and" por vírgula e transforma em lista
                df.at[index, 'sam'] = [item.strip() for item in df.at[index, 'sam'].replace('and', ',').split(',')]

            if df.at[index, 'acc_imp']:
                # Substitui "and" por vírgula e transforma em lista
                df.at[index, 'acc_imp'] = [item.strip() for item in df.at[index, 'acc_imp'].replace('and', ',').split(',')]
        return df
    df_split = process_paragraphs(df_accuracy)
    def merge_rows(df):
        merged_data = []

        for i in range(0, len(df), 2):
            if i + 1 < len(df): 
                paragrafo_1 = df.at[i, 'Name_sol']
                paragrafo_2 = df.at[i + 1, 'Name_sol']
                
                # Remover "Level:" no início do parágrafo para garantir que o conteúdo correto seja mostrado
                if paragrafo_1.startswith("Level:"):
                    paragrafo_1 = ""  # Se o parágrafo começar com "Level:", ignora esse parágrafo
                # Se paragrafo_1 estiver vazio, usa paragrafo_2
                if not paragrafo_1:
                    paragrafo = paragrafo_2
                else:
                    paragrafo = paragrafo_1 + " " + paragrafo_2
                # Preenche a coluna 'Level' com o valor da linha que contém "Level:"
                level = df.at[i, 'Level'] if df.at[i, 'Level'] else df.at[i + 1, 'Level']
                # Assumindo os valores não vazios de sam, acc_n, acc_imp, acc_imp_level
                sam = df.at[i, 'sam'] if df.at[i, 'sam'] else df.at[i + 1, 'sam']
                acc_n = df.at[i, 'acc_n'] if df.at[i, 'acc_n'] else df.at[i + 1, 'acc_n']
                acc_imp = df.at[i, 'acc_imp'] if df.at[i, 'acc_imp'] else df.at[i + 1, 'acc_imp']
                acc_imp_level = df.at[i, 'acc_imp_level'] if df.at[i, 'acc_imp_level'] else df.at[i + 1, 'acc_imp_level']
                
                merged_data.append([paragrafo, level, sam, acc_n, acc_imp, acc_imp_level])
            else:
                # Para a última linha que não tem par, apenas adiciona
                merged_data.append([df.at[i, 'Name_sol'], df.at[i, 'Level'], df.at[i, 'sam'],
                                    df.at[i, 'acc_n'], df.at[i, 'acc_imp'], df.at[i, 'acc_imp_level']])
        
        merged_df = pd.DataFrame(merged_data, columns=['Name_sol', 'Level', 'sam', 'acc_n', 'acc_imp', 'acc_imp_level'])
        
        return merged_df
    merged_df = merge_rows(df_split)
    # Função para processar e limpar a coluna 'Level'
    def clean_level(df):
        # Para cada linha, ajusta a coluna 'Level' para manter apenas o número percentual
        df['Level'] = df['Level'].apply(lambda x: re.sub(r'[^0-9]', '', str(x)) if isinstance(x, str) and '%' in x else x)
        return df
    df_clean=clean_level(merged_df)
    
    def update_acc_imp_level(df, imp_df):
        # Substitui os valores da coluna 'acc_imp_level' com base nos valores de 'imp_spec' de imp_df
        for index, row in df.iterrows():
            # Verifica se há componentes na coluna 'acc_imp' e se a coluna 'acc_imp_level' existe
            if isinstance(row['acc_imp'], list):  # Confirma que é uma lista
                acc_imp_levels = []
                for component in row['acc_imp']:
                    # Procura pelo 'imp_name' correspondente ao componente e obtém o 'imp_spec'
                    imp_spec_value = imp_df[imp_df['imp_name'] == component]['imp_spec']
                    if not imp_spec_value.empty:
                        # Converte o valor de 'imp_spec' para float e adiciona ao array
                        acc_imp_levels.append(float(imp_spec_value.iloc[0]))
                    else:
                        # Caso o componente não seja encontrado, adiciona um valor nulo
                        acc_imp_levels.append(None)
                # Atualiza a coluna 'acc_imp_level' com o array de valores encontrados
                df.at[index, 'acc_imp_level'] = acc_imp_levels
        
        return df
    df_clean2 = update_acc_imp_level(df_clean, imp_df)
    def ensure_list(value):
        if isinstance(value, list):
            return value
        elif value is not None:
            return [value]  # Converte valores não-lista para uma lista
        else:
            return []  # Converte valores None em listas vazias

    df_clean2['acc_imp_level'] = df_clean2['acc_imp_level'].apply(ensure_list)

    
    # Função para atualizar a coluna 'acc_sam_level' com base nos dados de 'sam_tc_df'
    def update_acc_sam_level(df, sam_tc_df):
        # Verificar se a coluna 'acc_sam_level' já existe, se não, criar
        if 'acc_sam_level' not in df.columns:
            df['acc_sam_level'] = None
        
        # Substitui os valores da coluna 'acc_sam_level' com base nos valores de 'sam_tc' de sam_tc_df
        for index, row in df.iterrows():
            # Verifica se a coluna 'sam' é uma lista
            if isinstance(row['sam'], list):  # Confirma que é uma lista
                acc_sam_levels = []  # Inicia uma lista para armazenar os valores de 'sam_tc'
                for component in row['sam']:
                    # Procura pelo 'sam_tc_name' correspondente ao componente e obtém o 'sam_tc'
                    sam_tc_value = sam_tc_df[sam_tc_df['Name'] == component]['sam_tc']
                    if not sam_tc_value.empty:
                        # Converte o valor de 'sam_tc' para float e adiciona à lista
                        acc_sam_levels.append(float(sam_tc_value.values[0]))  # Armazena como número (float)
                    else:
                        # Caso o componente não seja encontrado, adiciona um valor nulo
                        acc_sam_levels.append(None)
                # Atualiza a coluna 'acc_sam_level' com a lista de valores encontrados
                df.at[index, 'acc_sam_level'] = str(acc_sam_levels)  # Converte a lista em string para salvar no DataFrame
        
        return df
    df_clean3 = update_acc_sam_level(df_clean2, sam_tc_df)
    #df_clean3
    # Unir os valores de df_std e df_sam para preencher o campo 'Name'
    name_options = df_std['stds'].tolist() + df_sam['samples'].tolist()

    # Creating an empty DataFrame with name_options as the 'name' column
    columns = ['name', 'range', 'unit', 'Potency', 'CF', 'Stock_imp_mg_mL', 'w', 'V1', 'p1', 'V2', 'p2', 'V3']
    df_info = pd.DataFrame(columns=columns)

    # Creating the input fields for each name in name_options
    for name in name_options:
        # Adding a row for each name in the list
        with st.expander(f"Details for {name}"):
            range_value = st.number_input(f"Range for {name}", min_value=0, value=10 if 'SAM' in name else 1, key=f"range_{name}")
            unit_value = st.selectbox(f"Unit for {name}", ['%', 'mg', 'g'], key=f"unit_{name}")
            potency_value = st.number_input(f"Potency for {name}", min_value=0.0, max_value=1.0, value=1.0, key=f"potency_{name}")
            cf_value = st.number_input(f"Conversion Factor for {name}", min_value=0.0, value=1.0, key=f"cf_{name}")
            stock_concentration_value = st.number_input(f"Impurity Stock Concentration (mg/mL) for {name}", min_value=0.0, value=0.2 if 'IMP' in name else 0.0, key=f"stock_concentration_{name}")
            weight_value = st.number_input(f"Desired Weight (mg) for {name}", min_value=0, value=25 if 'SAM' in name else 20, key=f"weight_{name}")
            v1_value = st.number_input(f"Volume 1 (mL) for {name}", min_value=0, value=50 if 'SAM' in name else 100, key=f"v1_{name}")
            p1_value = st.number_input(f"Pipetted Volume 1 (mL) for {name}", min_value=0, value=5, key=f"p1_{name}")
            v2_value = st.number_input(f"Volume 2 (mL) for {name}", min_value=0, value=10, key=f"v2_{name}")
            p2_value = st.number_input(f"Pipetted Volume 2 (mL) for {name}", min_value=0, value=5, key=f"p2_{name}")
            v3_value = st.number_input(f"Volume 3 (mL) for {name}", min_value=0, value=5, key=f"v3_{name}")
            
            # Creating a new DataFrame with the entered values for this row
            new_row = pd.DataFrame({
                'name': [name],
                'range': [range_value],
                'unit': [unit_value],
                'Potency': [potency_value],
                'CF': [cf_value],
                'Stock_imp_mg_mL': [stock_concentration_value],
                'w': [weight_value],
                'V1': [v1_value],
                'p1': [p1_value],
                'V2': [v2_value],
                'p2': [p2_value],
                'V3': [v3_value]
            })
        
            # Concatenating the new row to the existing DataFrame
            df_info = pd.concat([df_info, new_row], ignore_index=True)

    # Creating a DataFrame with the collected inputs
    #df_info = pd.DataFrame(data_entries)

    # Display the final DataFrame with the entered data
    #st.write("The data you entered:")
    #st.dataframe(df_info)
    def add_tc_and_unit(row):
        name = row['name']
        
        # Verificar se o nome está em sam_tc_df
        if name in sam_tc_df['Name'].values:
            sam_tc = sam_tc_df[sam_tc_df['Name'] == name]['sam_tc'].values[0]
            sam_tc_unit = sam_tc_df[sam_tc_df['Name'] == name]['sam_tc_unit'].values[0]
            return pd.Series([sam_tc, sam_tc_unit], index=['tc', 'tc_unit'])
        
        # Verificar se o nome está em TC_std_df
        elif name in TC_std_df['TC_std_name'].values:
            tc_std = TC_std_df[TC_std_df['TC_std_name'] == name]['TC_std'].values[0]
            tc_std_unit = TC_std_df[TC_std_df['TC_std_name'] == name]['TC_std_unit'].values[0]
            return pd.Series([tc_std, tc_std_unit], index=['tc', 'tc_unit'])
        
        # Verificar se o nome está em imp_df
        elif name in imp_df['imp_name'].values:
            imp_spec = imp_df[imp_df['imp_name'] == name]['imp_spec'].values[0]
            imp_spec_unit = imp_df[imp_df['imp_name'] == name]['imp_spec_unit'].values[0]
            return pd.Series([imp_spec, imp_spec_unit], index=['tc', 'tc_unit'])
        
        # Se não encontrar o nome em nenhum DataFrame, retornar valores nulos
        return pd.Series([np.nan, np.nan], index=['tc', 'tc_unit'])

    # Aplicando a função para adicionar as colunas tc e tc_unit
    df_info[['tc', 'tc_unit']] = df_info.apply(add_tc_and_unit, axis=1)
    # Função para adicionar a coluna 'sam_tc' de acordo com as condições
    def add_sam_tc(row, sam_tc_value):
        if row['Category'] in ['sam', 'std']:
            return row['tc']  # Se a categoria for 'sam' ou 'std', sam_tc é igual a tc
        else:
            return sam_tc_value  # Caso contrário, sam_tc é igual ao primeiro valor de 'tc' onde a Category é 'sam'

    # Encontrar o primeiro valor de 'tc' onde 'Category' é 'sam'
    first_sam_tc = df_sel[df_sel['Category'] == 'sam'].iloc[0]['tc']

    # Aplicar a função para preencher a coluna 'sam_tc'
    df_sel['sam_tc'] = df_sel.apply(lambda row: add_sam_tc(row, first_sam_tc), axis=1)
    # Realizando a junção entre df_selt e df_info com base nas colunas 'Name' e 'name'
    df_sel = pd.merge(df_sel, df_info[['name', 'w','range','unit', 'Potency', 'CF', 'V1','V2','V3','p1', 'p2']], left_on='Name', right_on='name', how='left')

    # Remover a coluna 'name' após a junção, pois ela não é mais necessária
    df_sel.drop('name', axis=1, inplace=True)
    # Função para garantir que os valores sejam numéricos, convertendo strings em números
    def to_numeric(value):
        try:
            return pd.to_numeric(value, errors='coerce')  # 'coerce' converte valores inválidos para NaN
        except ValueError:
            return None

    # Função para calcular DF_sel para a categoria 'imp'
    def calculate_df_sel(row):
        if row['Category'] == 'imp':  
            sam_tc=row['sam_tc']
            tc=row['tc']
            w=row['w']

            df_sel_value=w*100/(tc*sam_tc)
            return df_sel_value
        else:
            return None  # Retorna None para outras categorias

    # Adicionar a coluna DF_sel a df_sel aplicando a função
    df_sel['DF_sel'] = df_sel.apply(calculate_df_sel, axis=1)

    # Filtrar VFs entre 10 e 250 e pipettes entre 0.5 e 10
    VF = [1, 2, 5, 10, 20, 25, 50, 100, 200, 250, 500, 1000, 2000]
    VF = [v for v in VF if 10 <= v <= 250]
    pipettes = [0.5, 1, 1.5, 2, 2.5, 3, 4, 5, 6, 7, 8, 9, 10, 12, 14, 15, 20, 25, 50]
    pipettes = [p for p in pipettes if 0.5 <= p <= 10]

    # Função para calcular o fator de diluição
    def calculate_dilution_factor(vf_comb, pipette_comb):
        vf_product = 1
        pipette_product = 1
        for vf in vf_comb:
            vf_product *= vf
            
        for pipette in pipette_comb:
            pipette_product *= pipette
        
        dilution_factor = vf_product / pipette_product
        return dilution_factor

    # Função para calcular a incerteza
    def calculate_uncertainty(vf_comb, pipette_comb):
        uncertainty = 0
        for i in range(len(pipette_comb)):
            uncertainty += abs(pipette_comb[i] - vf_comb[i + 1]) / vf_comb[i + 1]
        return uncertainty

    # Função para encontrar as combinações que resultam no fator de diluição desejado e minimizam a incerteza
    def find_best_dilution_combinations(target_factor):
        best_combination = None
        best_uncertainty = float('inf')
        
        for num_vfs in range(2, len(VF)+1):
            vf_combinations = itertools.combinations(VF, num_vfs)
            for vf_comb in vf_combinations:
                pipette_combinations = itertools.combinations(pipettes, num_vfs-1)
                for pipette_comb in pipette_combinations:
                    valid_combination = True
                    for i in range(len(pipette_comb)):
                        if pipette_comb[i] >= vf_comb[i + 1]:
                            valid_combination = False
                            break
                    if valid_combination:
                        dilution_factor = calculate_dilution_factor(vf_comb, pipette_comb)
                        if abs(dilution_factor - target_factor) / target_factor < 0.01:  # 1% de tolerância
                            uncertainty = calculate_uncertainty(vf_comb, pipette_comb)
                            if uncertainty < best_uncertainty:
                                best_uncertainty = uncertainty
                                best_combination = (vf_comb, pipette_comb)
        
        return best_combination

    # Função para ajustar as combinações de VF para cada linha
    def adjust_vf_combination(row, df_selt):
        target_factor = row['DF_sel']
        best_combination = None

        if row['Category'] in ['std', 'sam']:  # Para 'std' e 'sam'
            # Ajusta as combinações de VF para a categoria 'sam' conforme a lógica
            vf_comb = [row['V1'], row['V2'], row['V3']]  # Se existir
            pipette_comb = [row['p1'], row['p2']]  # Se existir
            
            # Se V3 for 1, use V2, se V2 for 1, use V1
            if row['V3'] == 1:
                vf_comb = [row['V2'], row['V1']]
            elif row['V2'] == 1:
                vf_comb = [row['V1']]
            
            # Garantir que o último valor seja sempre o maior VF
            vf_comb = [v for v in vf_comb if isinstance(v, (int, float))]  # Remove valores não numéricos
            vf_comb = tuple(sorted(vf_comb, reverse=True))  # Ordena em ordem decrescente para garantir o maior valor
            best_combination = (target_factor, vf_comb, tuple(pipette_comb))
            
        elif row['Category'] == 'imp':  # Para 'imp', escolha o maior VF (último valor)
            best_combination = find_best_dilution_combinations(target_factor)
            if best_combination:
                vf_comb, pipette_comb = best_combination
                vf_comb = [v for v in vf_comb if isinstance(v, (int, float))]  # Remove valores não numéricos
                vf_comb = tuple(sorted(vf_comb, reverse=True))  # Ordena em ordem decrescente para garantir o maior valor
                best_combination = (target_factor, vf_comb, tuple(pipette_comb))
            else:
                best_combination = (target_factor, None, None)

        # Garantir que sempre retorne uma tupla, mesmo que vazia
        if not best_combination:
            best_combination = (target_factor, (), ())

        return best_combination

    # Aplicar a função para cada linha do DataFrame
    best_combinations = []

    for index, row in df_sel.iterrows():
        best_combination = adjust_vf_combination(row, df_sel)
        best_combinations.append(best_combination)

    # Limpeza adicional para garantir que as tuplas não tenham valores vazios ('') e consistência com as combinações
    df_sel['Best_VF_combination'] = [comb[1] if comb[1] else () for comb in best_combinations]
    df_sel['Best_Pipette_combination'] = [comb[2] if comb[2] else () for comb in best_combinations]

    # Função para realizar as substituições conforme a descrição
    def update_columns_with_combinations(row, df):
        # Para categorias 'imp', 'std', 'sam', utilizar as combinações 'Best_VF_combination' e 'Best_Pipette_combination'
        if row['Category'] in ['imp', 'std', 'sam']:
            # Substituindo V1, V2, V3 com base nas tuplas de Best_VF_combination
            vf_combination = row['Best_VF_combination']
            p_combination = row['Best_Pipette_combination']
            
            # Para V1, V2, V3
            row['V1'] = vf_combination[0] if len(vf_combination) > 0 else ''
            row['V2'] = vf_combination[1] if len(vf_combination) > 1 else ''
            row['V3'] = vf_combination[2] if len(vf_combination) > 2 else ''
            
            # Para p1, p2
            row['p1'] = p_combination[0] if len(p_combination) > 0 else ''
            row['p2'] = p_combination[1] if len(p_combination) > 1 else ''
            
        # Para a categoria 'excip', substituir pelos valores onde app == Name
        elif row['Category'] == 'excip':
            app_name = row['app']
            # Buscar a linha correspondente onde Name é igual a app
            app_row = df[df['Name'] == app_name].iloc[0]
            
            # Atualizando as colunas V1, V2, V3, p1, p2 com os valores correspondentes
            row['V1'] = app_row['Best_VF_combination'][0] if len(app_row['Best_VF_combination']) > 0 else ''
            row['V2'] = app_row['Best_VF_combination'][1] if len(app_row['Best_VF_combination']) > 1 else ''
            row['V3'] = app_row['Best_VF_combination'][2] if len(app_row['Best_VF_combination']) > 2 else ''
            
            row['p1'] = app_row['Best_Pipette_combination'][0] if len(app_row['Best_Pipette_combination']) > 0 else ''
            row['p2'] = app_row['Best_Pipette_combination'][1] if len(app_row['Best_Pipette_combination']) > 1 else ''
            
        return row

    # Aplicando a função para cada linha do DataFrame
    df_sel = df_sel.apply(lambda row: update_columns_with_combinations(row, df_sel), axis=1)

    # Exibir as primeiras linhas do DataFrame para verificação
    #print(df_sel[['Name', 'V1', 'V2', 'V3', 'p1', 'p2']])
    
    def update_w_for_excip(row):
        # Verificar se a categoria é 'excip'
        if row['Category'] == 'excip' and pd.isnull(row['w']):
            # Garantir que V1, V2, V3, p1 e p2 são números válidos
            V1 = pd.to_numeric(row['V1'], errors='coerce') if pd.notnull(row['V1']) else 1
            V2 = pd.to_numeric(row['V2'], errors='coerce') if pd.notnull(row['V2']) else 1
            V3 = pd.to_numeric(row['V3'], errors='coerce') if pd.notnull(row['V3']) else 1
            p1 = pd.to_numeric(row['p1'], errors='coerce') if pd.notnull(row['p1']) else 1
            p2 = pd.to_numeric(row['p2'], errors='coerce') if pd.notnull(row['p2']) else 1
            
            # Se algum valor ainda for NaN após a conversão, substitui por 1
            V1 = V1 if not pd.isna(V1) else 1
            V2 = V2 if not pd.isna(V2) else 1
            V3 = V3 if not pd.isna(V3) else 1
            p1 = p1 if not pd.isna(p1) else 1
            p2 = p2 if not pd.isna(p2) else 1
            
            # Calcular w usando a fórmula fornecida
            row['w'] = (V1 * V2 * V3) / (p1 * p2) * row['tc']
            
        return row

    # Aplicar a função para atualizar os valores de 'w' para a categoria 'excip'
    df_selt = df_sel.apply(update_w_for_excip, axis=1)
    df_selt['w_tc'] = df_selt['w'] / (df_selt['Potency'] * df_selt['CF'])
    #st.title('Table for excip')
    #st.write(df_selt)
    # Função para calcular o comprimento total de 'tc', incluindo o ponto decimal
    def calculate_length_of_tc(row):
        # Conta todos os caracteres, incluindo o ponto decimal
        return len(str(row['tc']).replace('.', '')) + (1 if '.' in str(row['tc']) else 0)

    # Supondo que df_selt seja seu DataFrame
    df_selt['tc_length'] = df_selt.apply(calculate_length_of_tc, axis=1)

    # Função para calcular tc_min e tc_max com base no comprimento de tc
    def calculate_tc_min_max(row):
        # Calcula tc_min e tc_max com base na coluna 'tc' e no comprimento
        tc_min = round(row['tc'] - (4 / (10 ** (row['tc_length'] - (0 if row['tc_length'] == 1 else 1)))), row['tc_length'] + 1)
        tc_max = round(row['tc'] + (4 / (10 ** (row['tc_length'] - (0 if row['tc_length'] == 1 else 1)))), row['tc_length'] + 1)
        return pd.Series([tc_min, tc_max])  # Retorna como uma série para múltiplas colunas

    # Aplicando a função para calcular 'tc_min' e 'tc_max'
    df_selt[['tc_min', 'tc_max']] = df_selt.apply(calculate_tc_min_max, axis=1)
    # Função para calcular o número de casas decimais de um número
    def count_decimal_places(x):
        if isinstance(x, float):
            return len(str(x).split('.')[1]) if '.' in str(x) else 0
        return 0

    # Função para arredondar os valores de 'w_tc_min' e 'w_tc_max' para o mesmo número de casas decimais de 'tc'
    def round_to_tc_decimal_places(row):
        decimal_places = count_decimal_places(row['tc'])  # Contagem de casas decimais de 'tc'
        
        # Inicializando as variáveis
        w_tc_min = None
        w_tc_max = None
        
        # Condicional para Category = 'std' ou 'sam'
        if row['Category'] in ['std', 'sam']:
            # Se unit for 'mg', usa a lógica de w_tc ± range
            if row['unit'] == 'mg':
                w_tc_min = round(row['w'] - row['range'], 2)
                w_tc_max = round(row['w'] + row['range'], 2)
            # Se unit for '%', calcula w_tc ± (w_tc * range / 100)
            elif row['unit'] == '%':
                w_tc_min = round(row['w'] - (row['w_tc'] * row['range'] / 100), 2)
                w_tc_max = round(row['w'] + (row['w_tc'] * row['range'] / 100), 2)
        else:
            # Caso contrário, calcula como foi feito no código anterior
            w_tc_min = row['tc_min'] * row['w_tc'] / row['tc']
            w_tc_max = row['tc_max'] * row['w_tc'] / row['tc']
            
            # Arredonda para o número correto de casas decimais de 'tc'
            w_tc_min = round(w_tc_min, 2)
            w_tc_max = round(w_tc_max, 2)
        
        # Retorna as variáveis calculadas
        return pd.Series([w_tc_min, w_tc_max])
    #st.write(df_selt)
    # Aplicando a função para calcular e arredondar os valores de 'w_tc_min' e 'w_tc_max'
    df_selt[['w_tc_min', 'w_tc_max']] = df_selt.apply(round_to_tc_decimal_places, axis=1)
    # Função para calcular D1, D2 e D3 com base nas condições mencionadas
    def calculate_D_values(row, df_selt):
        # Inicializando as variáveis Potency e CF
        potency = row['Potency'] if pd.notnull(row['Potency']) else 1
        cf = row['CF'] if pd.notnull(row['CF']) else 1
        V1 = pd.to_numeric(row['V1'], errors='coerce') if pd.notnull(row['V1']) else 1
        V2 = pd.to_numeric(row['V2'], errors='coerce') if pd.notnull(row['V2']) else 1
        V3 = pd.to_numeric(row['V3'], errors='coerce') if pd.notnull(row['V3']) else 1
        p1 = pd.to_numeric(row['p1'], errors='coerce') if pd.notnull(row['p1']) else 1
        p2 = pd.to_numeric(row['p2'], errors='coerce') if pd.notnull(row['p2']) else 1
        # Atribuindo V1 (w_tc para 'imp' ou 'excip', ou w para 'std' ou 'sam')
        if row['Category'] in ['imp', 'excip']:
            weight = row['w_tc']  # Usar w_tc para 'imp' ou 'excip'
        else:  # Para 'std' ou 'sam'
            weight = row['w']
        
        # Calculando D1 = Potency * CF / V1
        D1 = (weight*potency * cf) / V1
        
        # Se a categoria for 'imp', multiplicar por 100 e dividir pelo primeiro valor de 'tc' com categoria 'sam'
        if row['Category'] == 'imp':
            sam_tc_value = df_selt[df_selt['Category'] == 'sam']['tc'].iloc[0] if not df_selt[df_selt['Category'] == 'sam'].empty else 1
            D1 = D1 * 100 / sam_tc_value
        
        # Calculando D2 e D3 de forma similar
        D2 = D1 * p1/V2  # Exemplo de cálculo para D2
        D3 = D2 * p2/V3 # Exemplo de cálculo para D3
        
        return pd.Series([D1, D2, D3])

    # Aplicando a função para calcular D1, D2, e D3 e adicionar ao DataFrame
    df_selt[['D1', 'D2', 'D3']] = df_selt.apply(calculate_D_values, axis=1, df_selt=df_selt)

    # Exibir o DataFrame atualizado
    #df_selt[['w','w_tc','w_tc_min', 'w_tc_max', 'Potency', 'CF', 'D1', 'D2', 'D3']]

    # Criar colunas ausentes com valor 1
    colunas_necessarias = ["V1", "V2", "V3", "p1", "p2", "p3"]
    for col in colunas_necessarias:
        if col not in df_selt.columns:
            df_selt[col] = 1  # Criar coluna com valor 1 caso não exista

    # Converter valores vazios e não numéricos para NaN e depois para float
    df_selt[colunas_necessarias] = df_selt[colunas_necessarias].replace("", pd.NA)
    df_selt[colunas_necessarias] = df_selt[colunas_necessarias].apply(pd.to_numeric, errors="coerce")

    # Substituir valores nulos de forma segura
    df_selt = df_selt.fillna(1)

    # Verificar se não há zeros no denominador
    if (df_selt["p1"] * df_selt["p2"]).eq(0).any():
        raise ValueError("Erro: O denominador contém zero, o que causaria uma divisão por zero.")

    # Calcular a nova coluna DF_final
    df_selt["DF_final"] = (df_selt["V1"] * df_selt["V2"] * df_selt["V3"]) / (df_selt["p1"] * df_selt["p2"])

    def add_tc_columns_from_spiking(row):
        # Verifica se 'tc_spiking' é uma lista e se não está vazia
        if isinstance(row['tc_spiking'], list) and row['tc_spiking']:
            # Percorre o array 'tc_spiking' e cria as novas colunas
            for i, tc_value in enumerate(row['tc_spiking']):
                # Cria o nome da coluna como 'stock_spiking_solution_{i+1}_tc'
                column_name = f"stock_spiking_solution_{i + 1}_tc"
                row[column_name] = tc_value  # Atribui o valor de 'tc_spiking' à nova coluna
        
        return row

    # Aplicar a função a cada linha do DataFrame
    df_selt2 = df_selt.apply(add_tc_columns_from_spiking, axis=1)
    # Função para arredondar D2 para o mesmo número de casas decimais de tc
    def round_to_tc_decimal_places(tc, D2):
        decimal_places = len(str(tc).split('.')[1]) if '.' in str(tc) else 0  # Conta o número de casas decimais em tc
        return round(D2, decimal_places)

    # Função para adicionar as colunas de spiking e "_name" dinamicamente
    def add_stock_spiking_columns(row, df_selt):
        # Verifica se a categoria é 'imp' ou se 'spiked' contém impurezas
        if row['Category'] == 'imp' or (isinstance(row['spiked'], list) and row['spiked']):
            spiking_values = []
            spiking_names = []
            
            # Processa todas as impurezas listadas em 'spiked'
            for i, impurity in enumerate(row['spiked']):
                column_name = f"stock_spiking_solution_{i + 1}"
                column_name_with_name = f"stock_spiking_solution_{i + 1}_name"
                column_name_tc = f"stock_spiking_solution_{i + 1}_tc"  # Nome da coluna dinâmica de tc
                
                # Encontra a linha correspondente no DataFrame
                matching_row = df_selt[df_selt['Name'] == impurity]

                if not matching_row.empty:
                    # Removendo `.iloc[0]` da coluna `row[column_name_tc]`, pois já é um valor único (float)
                    rounded_D2 = round_to_tc_decimal_places(row[column_name_tc], matching_row['D2'].iloc[0])
                    
                    # Usa a coluna f"stock_spiking_solution_{i + 1}_tc" para a comparação
                    if rounded_D2 <= 1.5 * row[column_name_tc]:  # Aqui também removi `.iloc[0]`
                        spiking_values.append("D1")
                    else:
                        spiking_values.append("D2")
                    spiking_names.append(matching_row['Name'].iloc[0])
                else:
                    spiking_values.append("D2")  # Caso a impureza não seja encontrada
                    spiking_names.append(None)   # Caso a impureza não seja encontrada
                
                # Adiciona as colunas dinâmicas ao DataFrame
                row[column_name] = spiking_values[i]
                row[column_name_with_name] = spiking_names[i]
        
        return row
    # Aplicando a função sem um limite fixo de colunas de spiking
    df_selt2 = df_selt2.apply(add_stock_spiking_columns, axis=1, df_selt=df_selt2)

    def add_concentration_column(row, df_selt):
        # Identifica o número máximo de colunas de spiking (vai até o número de impurezas na lista 'spiked')
        max_spiked_columns = len(row['spiked']) if isinstance(row['spiked'], list) else 0
        
        # Para cada spiking solution, busca a concentração correspondente
        for i in range(max_spiked_columns):  # Agora a função se adapta ao número de impurezas presentes
            column_name = f"stock_spiking_solution_{i + 1}"
            column_name_with_name = f"stock_spiking_solution_{i + 1}_name"
            column_name_conc = f"stock_spiking_solution_{i + 1}_name_conc"

            # Verifica se a coluna de spiking existe e não é nula
            if column_name in row and row[column_name] is not None:
                # Obtém o nome da impureza (da coluna '_name')
                impurity_name = row[column_name_with_name]
                
                # Se o nome da impureza não for nulo, busca a concentração correspondente
                if impurity_name:
                    # Encontra a linha onde o nome da impureza está na coluna 'Name'
                    matching_row = df_selt[df_selt['Name'] == impurity_name]
                    
                    # Verifica se a linha correspondente foi encontrada
                    if not matching_row.empty:
                        # Agora pegamos o valor da coluna que corresponde ao nome de impureza em stock_spiking_solution
                        column_for_value = row[column_name]  # O valor que está na coluna que será a busca para D2, D3, etc.
                        
                        # Verifica se a coluna indicada está presente e pega o valor da linha correspondente
                        if column_for_value in matching_row.columns:
                            conc_value = matching_row[column_for_value].iloc[0]
                            row[column_name_conc] = conc_value  # Atribui o valor de concentração na nova coluna
                        else:
                            row[column_name_conc] = None  # Se a coluna indicada não existir, coloca None
                    else:
                        row[column_name_conc] = None  # Se não encontrar, coloca None
                else:
                    row[column_name_conc] = None  # Se o nome da impureza for None, coloca None
        
        return row
    # Aplicar a função a cada linha do DataFrame
    df_selt2 = df_selt2.apply(add_concentration_column, axis=1, df_selt=df_selt2)

    def add_volume_columns(row):
        # Verifica se 'tc_spiking' é uma lista e se as colunas necessárias existem
        if isinstance(row['tc_spiking'], list) and row['tc_spiking']:
            for i in range(len(row['tc_spiking'])):
                # Definir os nomes das colunas com base no índice
                tc_column_name = f"stock_spiking_solution_{i + 1}_tc"
                conc_column_name = f"stock_spiking_solution_{i + 1}_name_conc"
                volume_column_name = f"stock_spiking_solution_{i + 1}_volume"

                # Verifica se as colunas existem no DataFrame e calcula o valor
                if tc_column_name in row and conc_column_name in row:
                    tc_value = row[tc_column_name]
                    conc_value = row[conc_column_name]
                    df_final_value = row['DF_final']  # Supondo que a coluna 'DF_final' esteja presente

                    # Realiza o cálculo e atribui o valor à nova coluna
                    row[volume_column_name] = (tc_value * df_final_value) / conc_value
                else:
                    row[volume_column_name] = None  # Se as colunas não existirem, atribui None

        return row

    df_selt2 = df_selt2.apply(add_volume_columns, axis=1)
    #st.write(df_selt2)

    # Função para gerar o parágrafo baseado na categoria
    def generate_paragraph_text(row):
        paragraph = f"{row['Name_sol']}:\n"  # Inicia com o nome completo da linha
        
        # Para as categorias "", "std", "RES", ou "sam" com spiked vazio
        if row['Category'] in ["BLK", "std", "RES"] or (row['Category'] == "sam" and (not row['spiked'] or row['spiked'] == [""])):
            paragraph += "Prepare as per analytical method."
        
        # Para a categoria 'imp'
        elif row['Category'] == 'imp':  
            if pd.notnull(row['D2']):
                paragraph += f"\nStock I for {row['Name']} Solution:\nWeight {round(row['w_tc'], 2)} mg ({row['w_tc_min']} mg to {row['w_tc_max']} mg) into {row['V1']} mL and complete to volume."
            else:
                paragraph += f"\n\nWeight {round(row['w_tc'], 2)} mg ({row['w_tc_min']} mg to {row['w_tc_max']} mg) into {row['V1']} mL and complete to volume."
            
            # Adiciona diluições
            if pd.notnull(row['D3']):
                paragraph += f"\n\nStock II for {row['Name']} Solution:\nDilute {row['p1']} mL from Stock I into {row['V2']} mL volumetric flask and complete to volume."
            else:
                paragraph += f"\n\n{row['Name']} at {row['tc']}% solution:\nDilute {row['p1']} mL from Stock I into {row['V2']} mL volumetric flask and complete to volume."
            if pd.notnull(row['D3']):
                paragraph += f"\n\n{row['Name']} at {row['tc']}% solution:\nDilute {row['p2']} mL from Stock II into {row['V3']} mL volumetric flask and complete to volume."
        
        # Para a categoria 'excip'
        elif row['Category'] == 'excip':
            if pd.notnull(row['V2']) and row['V2'] == 1 and row['p1'] == 1:
                paragraph += f"Weight {round(row['w'], 2)} ± 10% into {row['V1']} mL volumetric flask and complete to volume."
            else:
                paragraph += f"Stock I for {row['Name']} Solution:\nWeight {round(row['w'], 2)} ± 10% into {row['V1']} mL volumetric flask and complete to volume."
                if pd.notnull(row['V3']) and row['V3'] == 1 and row['p2'] == 1:
                    paragraph += f"\nDilute {row['p2']} mL from Stock I into {row['V2']} mL volumetric flask and complete to volume."
                else:
                    paragraph += f"\nStock II for {row['Name']} Solution:\nDilute {row['p1']} mL from Stock I into {row['V2']} mL volumetric flask and complete to volume."
        
        # Para a categoria 'sam' com impurezas
        elif row['Category'] == 'sam' and row['spiked'] != ['']:
            stock_spiking_phrases = []  # Lista para armazenar os volumes de spiking

            # Percorre todas as colunas dinâmicas de spiking
            for i in range(1, 100):  # Assume até 100 spikings para cobrir todas as colunas
                col_name = f"stock_spiking_solution_{i}_name"
                col_volume = f"stock_spiking_solution_{i}_volume"
                col_type = f"stock_spiking_solution_{i}"  # Verifica se é D1 ou D2
                
                if col_name in row and pd.notnull(row[col_name]):  # Se a impureza existir
                    solution_type = "Stock I" if row[col_type] == "D1" else "Stock II"
                    stock_spiking_phrases.append(f"{row[col_volume]} mL of {solution_type} solution for {row[col_name]}")

            # Construção correta da frase dinâmica
            stock_spiking_text = ""
            if stock_spiking_phrases:
                if len(stock_spiking_phrases) == 1:
                    stock_spiking_text = stock_spiking_phrases[0]
                else:
                    stock_spiking_text = ", ".join(stock_spiking_phrases[:-1]) + " and " + stock_spiking_phrases[-1]

            # Construção do texto principal
            if pd.notnull(row['V2']) and row['V2'] == 1 and row['p1'] == 1:
                paragraph += f"\nWeight {round(row['w'], 2)} ± 10% into {row['V1']} mL volumetric flask"
                if stock_spiking_text:
                    paragraph += f", add {stock_spiking_text}"
                paragraph += " and complete to volume."
            else:  
                paragraph += f"Stock I for {row['Name']} spiked solution:\nWeight {round(row['w'], 2)} ± 10% into {row['V1']} mL volumetric flask"
                if stock_spiking_text:
                    paragraph += f", add {stock_spiking_text}"
                paragraph += " and complete to volume."      

                # Construção da parte de Stock II
                if pd.notnull(row['V3']) and row['V3'] != 1 and row['p2'] != 1:
                    paragraph += f"\nDilute {row['p2']} mL from Stock I into {row['V2']} mL volumetric flask and complete to volume."
                else:
                    paragraph += f"\nStock II for {row['Name']} Solution:\nDilute {row['p1']} mL from Stock I into {row['V2']} mL volumetric flask and complete to volume."

        return paragraph

    # Gerar os parágrafos para cada linha e salvar como texto
    paragraphs = df_selt2.apply(generate_paragraph_text, axis=1)

    # Exibir os parágrafos no Streamlit
    st.title("Generated Paragraphs for Method Validation")

    # Usando um loop para exibir os parágrafos gerados na interface do Streamlit
    for paragraph in paragraphs:
        st.text(paragraph + "\n")

else:
    st.warning("Por favor, carregue um documento DOCX.")
