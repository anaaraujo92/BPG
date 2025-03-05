import streamlit as st
import pandas as pd
import zipfile
from fpdf import FPDF
from docx import Document
from lxml import etree
import io
import re
import itertools
import numpy as np

def extract_content_controls(docx_path, target_tags):
    '''function to exctracts the content control text associated to specific tags'''
    with zipfile.ZipFile(docx_path, 'r') as docx_zip:
        xml_content = docx_zip.read('word/document.xml')
    root = etree.XML(xml_content)
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    extracted_data = {tag: [] for tag in target_tags}
    for sdt in root.xpath('.//w:sdt', namespaces=ns):
        tag_element = sdt.find('.//w:tag', namespaces=ns)
        text_element = sdt.find('.//w:t', namespaces=ns)
        if tag_element is not None and text_element is not None:
            tag_value = tag_element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
            if tag_value in target_tags:
                extracted_data[tag_value].append(text_element.text)
    return extracted_data

def extract_specific_table(docx_file, search_text):
    '''function to exctracts the content control text associated to tables'''
    doc = Document(docx_file)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if search_text.lower() in cell.text.strip().lower():  
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        table_data.append(row_data)
                    return table_data 
    return None        


st.title("Bench Plan Generator for Method Validation")

# File upload
docx_file = st.file_uploader("Upload Word Document", type=["docx"])

if docx_file:
    # Save the uploaded file
    file_path = f"/tmp/{docx_file.name}"
    with open(file_path, "wb") as f:
        f.write(docx_file.getbuffer())
    # Extracting 'std' data
    target_tags = ["std", "std_type"]
    std_data = extract_content_controls(docx_file, target_tags)
    df_std = pd.DataFrame({"stds": std_data["std"],
                        "std_type": std_data["std_type"]})
    # Extracting 'sam' data
    target_tags = ["sam"]
    sam_data = extract_content_controls(docx_file, target_tags)
    df_sam = pd.DataFrame(sam_data['sam'], columns=["samples"])
    # Extracting 'excip' data
    target_tags = ["excip"]
    excip_data = extract_content_controls(docx_file, target_tags)
    df_excip = pd.DataFrame(excip_data['excip'], columns=["sample"])
    # Extracting composition data
    target_tags = ["composition_name", "composition", "composition_per"]
    composition_data = extract_content_controls(docx_file, target_tags)
    min_length = min(len(composition_data["composition_name"]), len(composition_data["composition"]), len(composition_data["composition_per"]))
    df_composition = pd.DataFrame({
        "sample": composition_data["composition_name"][:min_length],
        "composition": composition_data["composition"][:min_length],
        "ratio": composition_data["composition_per"][:min_length]
    })
    target_tags = ["TC_std_name", "TC_std_level", "TC_std_relation", "TC_std_unit"]
    TC_std_data = extract_content_controls(docx_file, target_tags)
    TC_std_df = pd.DataFrame({
        "TC_std_name": TC_std_data["TC_std_name"],
        "TC_std_level": TC_std_data["TC_std_level"],
        "TC_std_relation": TC_std_data["TC_std_relation"],
        "TC_std_unit": TC_std_data["TC_std_unit"]
    })
    search_text = "in relation to"
    TC_std_df_w = extract_specific_table(docx_file, search_text)

    if TC_std_df_w:
        for row in TC_std_df_w:
            print(row)
    else:
        print("Table not found.")
    df = pd.DataFrame(TC_std_df_w[1:], columns=TC_std_df_w[0])
    df.columns = ['Name', 'TC_std', 'Concentration']
    TC_std_df_w = df[['TC_std']]
    TC_std_df = pd.concat([TC_std_df, TC_std_df_w], axis=1)
    def to_numeric(value):
        try:
            return pd.to_numeric(value, errors='coerce')  # to force to convert to numeric
        except ValueError:
            return None
    TC_std_df['TC_std'] = pd.to_numeric(TC_std_df['TC_std'], errors='coerce')
    target_tags = ["imp_spec_unit", "imp_loq_unit", "imp_lod_unit"]
    imp_data = extract_content_controls(docx_file, target_tags)
    imp_df = pd.DataFrame({
        "imp_spec_unit": imp_data["imp_spec_unit"],
        "imp_loq_unit": imp_data["imp_loq_unit"],
        "imp_lod_unit": imp_data["imp_lod_unit"]
    })
    search_text = "Impurities Specification/limit" 
    imp_df_w = extract_specific_table(docx_file, search_text)
    if imp_df_w:
        for row in imp_df_w:
            print(row)
    else:
        print("Table not found.")
    df = pd.DataFrame(imp_df_w[1:], columns=imp_df_w[0])
    df.columns = ['imp_name', 'imp_spec', 'imp_spec_unit', 'imp_loq', 'imp_loq_unit', 'imp_lod', 'imp_lod_unit']
    imp_df_w = df[['imp_name','imp_spec', 'imp_loq', 'imp_lod']]
    imp_df = pd.concat([imp_df_w, imp_df], axis=1)
    imp_df['imp_spec'] = pd.to_numeric(imp_df['imp_spec'], errors='coerce')
    imp_df['imp_loq'] = pd.to_numeric(imp_df['imp_loq'], errors='coerce')
    imp_df['imp_lod'] = pd.to_numeric(imp_df['imp_lod'], errors='coerce')
    imp_df.info()
    target_tags = ["imp_spec_u_unit"]

    imp_df_u_data = extract_content_controls(docx_file, target_tags)
    imp_df_u = pd.DataFrame({
        "imp_spec_u_unit": imp_df_u_data["imp_spec_u_unit"]
    })
    search_text = "Any unknown related substances" 
    imp_df_u_w = extract_specific_table(docx_file, search_text)
    df = pd.DataFrame(imp_df_u_w[1:], columns=imp_df_u_w[0])
    df.columns = ['Name', 'imp_spec', 'imp_spec_unit']
    imp_df_u_w = df[['Name', 'imp_spec']]
    imp_df_u = pd.concat([imp_df_u, imp_df_u_w], axis=1)
    imp_df_u['imp_spec'] = pd.to_numeric(imp_df_u['imp_spec'], errors='coerce')
    imp_df_u.info()
    target_tags = ["sam_tc_unit", "sam_tc_app"]
    sam_tc_df_data = extract_content_controls(docx_file, target_tags)
    sam_tc_df = pd.DataFrame({
        "sam_tc_unit": sam_tc_df_data["sam_tc_unit"],
        "sam_tc_app": sam_tc_df_data["sam_tc_app"]
    })
    search_text = "Sample Concentration" 
    sam_tc_df_w = extract_specific_table(docx_file, search_text)
    df = pd.DataFrame(sam_tc_df_w[1:], columns=sam_tc_df_w[0])
    df.columns = ['Name', 'sam_tc', 'sam_tc_unit', 'sam_tc_app']
    sam_tc_df_w = df[['Name', 'sam_tc']]
    sam_tc_df = pd.concat([sam_tc_df_w, sam_tc_df], axis=1)
    sam_tc_df['sam_tc'] = pd.to_numeric(sam_tc_df['sam_tc'], errors='coerce')
    target_tags = ["excip_tc_unit"]
    excip_tc_df_data = extract_content_controls(docx_file, target_tags)
    excip_tc_df = pd.DataFrame({
        "excip_tc_unit": excip_tc_df_data["excip_tc_unit"]
    })
    search_text = "Excipients/Placebo Concentration" 
    excip_tc_df_w = extract_specific_table(docx_file, search_text)
    df = pd.DataFrame(excip_tc_df_w[1:], columns=excip_tc_df_w[0])
    df.columns = ['Name', 'excip_tc', 'excip_tc_unit', 'df_sam']
    excip_tc_df_w = df[['Name', 'excip_tc','df_sam']]
    excip_tc_df = pd.concat([excip_tc_df_w, excip_tc_df], axis=1)
    excip_tc_df['excip_tc'] = pd.to_numeric(excip_tc_df['excip_tc'], errors='coerce')
    
    def extraction_parameter(docx_file, word):
        """
        Function to find a specific word in a cell and divide the content in paragraphs.
        
        :param docx_file: document file (uploaded as BytesIO object)
        :param word: word to be found
        
        :return: DataFrame with extracted paragraphs.
        """
        if not docx_file or not isinstance(docx_file, io.BytesIO):
            return None  # Ensure the file is valid

        try:
            # Use io.BytesIO to load file content
            doc = Document(io.BytesIO(docx_file.getvalue()))  
        except Exception as e:
            print(f"Error processing DOCX: {e}")
            return None

        dados = []
        for tabela in doc.tables:
            for linha in tabela.rows:
                for i, celula in enumerate(linha.cells):
                    texto_celula = celula.text.strip()
                    if word in texto_celula:
                        if i + 1 < len(linha.cells):
                            texto_direita = linha.cells[i + 1].text.strip()
                            paragrafos = texto_direita.split("\n")
                            for par in paragrafos:
                                dados.append([par.strip()])
        df = pd.DataFrame(dados, columns=["Name_sol"])
        df = df.dropna()
        return df

    # Extract 'Selectivity' Data
    df_selectivity = extraction_parameter(docx_file, 'Selectivity')
    if df_selectivity is not None and not df_selectivity.empty:
        st.write("Selectivity data extracted")
    else:
        st.warning("Error processing the file or 'Selectivity' not found.")
    
    # Drop the first index row and reset the index
    if df_selectivity is not None and not df_selectivity.empty:
        # Strip leading/trailing spaces from the 'Name_sol' column and filter out rows that are empty
        df_selectivity = df_selectivity[df_selectivity["Name_sol"].str.strip() != ""]
        # Drop the first row and reset the index
        df_sel = df_selectivity.drop(index=0).reset_index(drop=True)
        st.write("Selectivity data processed")
    else:
        st.warning("Selectivity data could not be extracted or is empty.")

    # Extracting selectivity content into columns
    def extract_sel(df):
        """
        Function to divide selectivity content into columns: Name, tc, app, spiked, and tc_spiking.
        """
        def slipt_rows(linha):
            name_match = re.match(r"^(.*?)\sat", linha)
            tc_match = re.search(r"(at\s(?:method|sample|specification)\s.*\sconcentration)", linha)
            app_match = re.search(r"for\s(.*)", linha)
            spiked_match = re.search(r"spiked\s(.*?)\sat", linha)
            # 'tc_spiking' - The remaining part of the sentence after the last extraction
            tc_spiking_match = re.search(r"at\s.*\s(.*)", linha)
            # Extraction of the parts, if they exist
            name = name_match.group(1) if name_match else linha
            tc = tc_match.group(1) if tc_match else ""
            app = app_match.group(1) if app_match else ""
            spiked = spiked_match.group(1) if spiked_match else ""
            # Additional Verification to avoid error in case tc_spiking_match is None
            tc_spiking = tc_spiking_match.group(1) if tc_spiking_match and tc_spiking_match.group(1) else ""
            return name, tc, app, spiked, tc_spiking
        df[['Name', 'tc', 'app', 'spiked', 'tc_spiking']] = df['Name_sol'].apply(lambda x: pd.Series(slipt_rows(x)))
        return df

    df_sel = extract_sel(df_sel)
    df_sel.loc[~df_sel['Name_sol'].str.contains('spiked', case=False), 'tc_spiking'] = None
    def assign_column_value(row):
        if row['Name'] == 'Standard':
            return 'std'
        elif row['Name'] in imp_df['imp_name'].values:
            return 'imp'
        elif row['Name'] in df_sam['samples'].values:
            return 'sam'
        elif row['Name'] in df_excip['sample'].values:
            return 'excip'
        elif row['Name'] == 'Resolution':
            return 'RES'
        elif 'blank' in row['Name'].lower():  # Check if 'blank' is in the Name (case insensitive)
            return 'BLK'
        else:
            return None
    
    df_sel['Category'] = df_sel.apply(assign_column_value, axis=1)
    def assign_column_value_app(row):
        if row['app'] == 'Standard':
            return 'std'
        elif row['app'] in imp_df['imp_name'].values:
            return 'imp'
        elif row['app'] in df_sam['samples'].values:
            return 'sam'
        elif row['app'] in df_excip['sample'].values:
            return 'excip'
        elif row['app'] == 'Resolution':
            return 'RES'
        elif row['app'] == '':
            return 'BLK'
        else:
            return None 
    df_sel['Category_app'] = df_sel.apply(assign_column_value_app, axis=1)
    def update_tc(row):
        if row['Category'] == 'std':
            tc_value = TC_std_df[~TC_std_df['TC_std_level'].isin(['LOQ', 'LOD'])]['TC_std']
            if not tc_value.empty:
                return tc_value.iloc[0]
            else:
                return None  
        elif row['Category'] == 'imp':
            tc_value = imp_df[imp_df['imp_name'] == row['Name']]['imp_spec']
            if not tc_value.empty:
                return tc_value.iloc[0]
            else:
                return None
        elif row['Category'] == 'excip':
            tc_value = excip_tc_df[(excip_tc_df['Name'] == row['Name']) & 
                                    (excip_tc_df['df_sam'] == row['app'])]['excip_tc']
            if not tc_value.empty:
                return tc_value.iloc[0]
            else:
                return None
        elif row['Category'] == 'sam':
            tc_value = sam_tc_df[sam_tc_df['Name'] == row['Name']]['sam_tc']
            if not tc_value.empty:
                return tc_value.iloc[0]
            else:
                return None
        return None
    df_sel['tc'] = df_sel.apply(update_tc, axis=1)
    def modify_spiked(spiked_value):
        modified_value = spiked_value.replace('with ', '').replace('and ', ', ')
        return [item.strip() for item in modified_value.split(',')]
    df_sel['spiked'] = df_sel['spiked'].apply(modify_spiked)
    def modify_tc_spiking(row):
        if row['tc_spiking'] == 'specification':
            values = []
            for item in row['spiked']:
                imp_spec_value = imp_df.loc[imp_df['imp_name'] == item, 'imp_spec']
                if not imp_spec_value.empty:
                    values.append(float(imp_spec_value.iloc[0]))
                else:
                    values.append(None)
            return values 
        elif isinstance(row['tc_spiking'], str) and '%' in row['tc_spiking']:
            try:
                value = float(row['tc_spiking'].replace('%', ''))
                return [value] * len(row['spiked'])
            except ValueError:
                return [None] * len(row['spiked'])
        elif pd.notnull(row['tc_spiking']) and row['tc_spiking'] != '':
            try:
                return [float(row['tc_spiking'])] * len(row['spiked'])
            except ValueError:
                return [None] * len(row['spiked'])
        return [None] * len(row['spiked'])
    df_sel['tc_spiking'] = df_sel.apply(modify_tc_spiking, axis=1)
    def extraction_accuracy(doc_path):
        """
        Function to locate 'Accuracy' and extract paragraohs tant contains "For known" and "For Unknwon" without include them.
        :param doc_path: document path.
        :return: DataFrame with paragraohs.
        """
        if not docx_file or not isinstance(docx_file, io.BytesIO):
            return None  # Ensure the file is valid
        try:
            # Use io.BytesIO to load file content
            doc = Document(io.BytesIO(docx_file.getvalue()))  
        except Exception as e:
            print(f"Error processing DOCX: {e}")
            return None
        dados = []
        extracting = False
        for tabela in doc.tables:
            for linha in tabela.rows:
                for i, celula in enumerate(linha.cells):
                    texto_celula = celula.text.strip()
                    if "Accuracy" in texto_celula:
                        print(f"'Accuracy' found in cell: {texto_celula}")
                        if i + 1 < len(linha.cells):  
                            texto_direita = linha.cells[i + 1].text.strip()
                            print(f"Content in cell at rigth: {texto_direita}")
                            paragrafos = texto_direita.split("\n")
                            for par in paragrafos:
                                par = par.strip()
                                if "for known" in par.lower(): 
                                    extracting = True
                                    print(f"Extraction starting at: {par}")
                                if "for unknown" in par.lower():
                                    extracting = False
                                    print(f"Extraction stopping at: {par}")
                                    break
                                if extracting:
                                    print(f"Adding pragraph: {par}")
                                    dados.append([par])
                        else:
                            print("No cell found.")
        if not dados:
            print("No data extracted.")
        else:
            print(f"{len(dados)} Paragraph extracted.")

        df = pd.DataFrame(dados, columns=["Name_sol"])
        df = df.dropna()
        
        return df
      
    # Extract 'Selectivity' Data
    df_accuracy = extraction_accuracy(docx_file)
    if df_accuracy is  not None and not df_accuracy.empty:
        st.write("Extracted Accuracy Data")  # Display DataFrame if it contains data
    else:
        st.warning("Error processing the file or 'Selectivity' not found.")
    df_accuracy = df_accuracy[df_accuracy["Name_sol"].str.strip() != ""]
    # Drop the first index row and reset
    if df_accuracy is not None and not df_accuracy.empty:
        df_accuracy = df_accuracy.drop(index=0).reset_index(drop=True)
        st.write("Processed Accuracy Data")  # Show processed selectivity data
    else:
        st.warning("Acccuracy data could not be processed.")
    #df_accuracy = pd.DataFrame(df_accuracy)
    def process_paragraphs(df):
        df['Level'] = None
        df['sam'] = None
        df['acc_n'] = None
        df['acc_imp'] = None
        df['acc_imp_level'] = None
        for index, row in df.iterrows():
            texto = row['Name_sol']
            if "Level:" in texto:
                df.at[index, 'Level'] = texto.split(":")[1].strip()
            # Filling column 'sam' from 'For' to 'sample (s)'
            if texto.startswith("For") and "samples (s)," in texto:
                sam_match = re.search(r"^For\s+(.*?)\s+samples \(s\),", texto)
                if sam_match:
                    df.at[index, 'sam'] = sam_match.group(1)
            # Filling columna 'acc_imp' from "spiked with" to "at", without including them
            if "spiked with" in texto and "at" in texto:
                acc_imp_match = re.search(r"spiked with(.*?)at", texto)
                if acc_imp_match:
                    df.at[index, 'acc_imp'] = acc_imp_match.group(1).strip()
            # Filling column 'acc_imp_level' from "at" until the end
            if "at" in texto:
                acc_imp_level_match = re.search(r"at(.*)", texto)
                if acc_imp_level_match:
                    df.at[index, 'acc_imp_level'] = acc_imp_level_match.group(1).strip()
            # Filling column 'acc_n' with the following conditions
            if "Prepare" in texto and "as is" in texto:
                # When "as is", extract from "Prepare" until the first ")"
                acc_n_match = re.search(r"Prepare(.*?[\)])", texto)
                if acc_n_match:
                    df.at[index, 'acc_n'] = acc_n_match.group(1).strip()
            elif "prepare" in texto:
                # when "Prepare", extract from "Prepare" to "independent", without including them"
                acc_n_match = re.search(r"prepare(.*?)independent", texto, re.IGNORECASE)
                if acc_n_match:
                    df.at[index, 'acc_n'] = acc_n_match.group(1).strip()
            # Filling column 'sam' with the follwing conditions
            if "as is" in texto:
                # Extrating from the first ")" to"sample", without including them
                sam_match = re.search(r"\)(.*?)sample", texto)
                if sam_match:
                    df.at[index, 'sam'] = sam_match.group(1).strip()
            # if the text don't include in the above conditions, the paragraph is all included.
            if pd.isna(df.at[index, 'Level']) and pd.isna(df.at[index, 'sam']) and pd.isna(df.at[index, 'acc_n']) and pd.isna(df.at[index, 'acc_imp']) and pd.isna(df.at[index, 'acc_imp_level']):
                df.at[index, 'sam'] = texto
            # For column acc_n, extracting only the number within ()
            if df.at[index, 'acc_n']:
                acc_n_number = re.search(r"\((\d+)\)", df.at[index, 'acc_n'])
                if acc_n_number:
                    df.at[index, 'acc_n'] = acc_n_number.group(1) 
            # For columns 'sam' and 'acc_imp', replace 'and' for ',' and transfor into list
            if df.at[index, 'sam']:
                df.at[index, 'sam'] = [item.strip() for item in df.at[index, 'sam'].replace('and', ',').split(',')]
            if df.at[index, 'acc_imp']:
                df.at[index, 'acc_imp'] = [item.strip() for item in df.at[index, 'acc_imp'].replace('and', ',').split(',')]
        return df
    df_split = process_paragraphs(df_accuracy)
    def merge_rows(df):
        '''Merges consecutive rows in the given DataFrame by concatenating relevant fields.
         The function ensures that the 'Level', 'sam', 'acc_n', 'acc_imp', and 'acc_imp_level' 
        fields are taken from the first non-empty occurrence in the pair. If there is an 
        odd number of rows, the last row is added as is.'''
        merged_data = []
        for i in range(0, len(df), 2):
            if i + 1 < len(df): 
                paragrafo_1 = df.at[i, 'Name_sol']
                paragrafo_2 = df.at[i + 1, 'Name_sol']
                # Remove "Level:" 
                if paragrafo_1.startswith("Level:"):
                    paragrafo_1 = ""
                if not paragrafo_1:
                    paragrafo = paragrafo_2
                else:
                    paragrafo = paragrafo_1 + " " + paragrafo_2
                level = df.at[i, 'Level'] if df.at[i, 'Level'] else df.at[i + 1, 'Level']
                sam = df.at[i, 'sam'] if df.at[i, 'sam'] else df.at[i + 1, 'sam']
                acc_n = df.at[i, 'acc_n'] if df.at[i, 'acc_n'] else df.at[i + 1, 'acc_n']
                acc_imp = df.at[i, 'acc_imp'] if df.at[i, 'acc_imp'] else df.at[i + 1, 'acc_imp']
                acc_imp_level = df.at[i, 'acc_imp_level'] if df.at[i, 'acc_imp_level'] else df.at[i + 1, 'acc_imp_level']
                merged_data.append([paragrafo, level, sam, acc_n, acc_imp, acc_imp_level])
            else:
                # For the last row
                merged_data.append([df.at[i, 'Name_sol'], df.at[i, 'Level'], df.at[i, 'sam'],
                                    df.at[i, 'acc_n'], df.at[i, 'acc_imp'], df.at[i, 'acc_imp_level']])
        merged_df = pd.DataFrame(merged_data, columns=['Name_sol', 'Level', 'sam', 'acc_n', 'acc_imp', 'acc_imp_level'])
        return merged_df
    merged_df = merge_rows(df_split)
    def clean_level(df):
        '''Function to process and clena column Level, by keeping only the % value'''
        df['Level'] = df['Level'].apply(lambda x: re.sub(r'[^0-9]', '', str(x)) if isinstance(x, str) and '%' in x else x)
        return df
    df_clean=clean_level(merged_df)
    def update_acc_imp_level(df, imp_df):
        # Replacing values of column acc_imp_level based on 'imp_spec' from imp_df
        for index, row in df.iterrows():
            if isinstance(row['acc_imp'], list):  # List confirmation
                acc_imp_levels = []
                for component in row['acc_imp']:
                    # Look for 'imp_name' and obtains 'imp_spec'
                    imp_spec_value = imp_df[imp_df['imp_name'] == component]['imp_spec']
                    if not imp_spec_value.empty:
                        acc_imp_levels.append(float(imp_spec_value.iloc[0]))
                    else:
                        acc_imp_levels.append(None)
                df.at[index, 'acc_imp_level'] = acc_imp_levels
        return df
    df_clean2 = update_acc_imp_level(df_clean, imp_df)
    def ensure_list(value):
        if isinstance(value, list):
            return value
        elif value is not None:
            return [value] 
        else:
            return []  
    df_clean2['acc_imp_level'] = df_clean2['acc_imp_level'].apply(ensure_list)
    
    def update_acc_sam_level(df, sam_tc_df):
        '''Function to update column acc_sam_level based on sam_tc_df'''
        if 'acc_sam_level' not in df.columns:
            df['acc_sam_level'] = None
        for index, row in df.iterrows():
            if isinstance(row['sam'], list): 
                acc_sam_levels = [] 
                for component in row['sam']:
                    sam_tc_value = sam_tc_df[sam_tc_df['Name'] == component]['sam_tc']
                    if not sam_tc_value.empty:
                        acc_sam_levels.append(float(sam_tc_value.values[0])) 
                    else:
                        acc_sam_levels.append(None)
                df.at[index, 'acc_sam_level'] = str(acc_sam_levels) 
        return df
    df_clean3 = update_acc_sam_level(df_clean2, sam_tc_df)
    # List containing all standards/impurities and samples
    name_options = df_std['stds'].tolist() + df_sam['samples'].tolist()
    columns = ['name', 'range', 'unit', 'Potency', 'CF', 'Stock_imp_mg_mL', 'w', 'V1', 'p1', 'V2', 'p2', 'V3']
    df_info = pd.DataFrame(columns=columns)
    # Creating the input fields for each name in name_options
    for name in name_options:
        # Adding a row for each name in the list
        with st.expander(f"Details for {name}"):
            range_value = st.number_input(f"Range for {name}", min_value=0, value=10 if 'SAM' in name else 1, key=f"range_{name}")
            unit_value = st.selectbox(f"Unit for {name}", ['%', 'mg', 'g'], key=f"unit_{name}")
            potency_value = st.number_input(f"Potency for {name}", min_value=0.0, max_value=1.0, value=1.0, key=f"potency_{name}")
            cf_value = st.number_input(f"Conversion Factor for {name}", min_value=0.0, value=1.0, key=f"cf_{name}")
            stock_concentration_value = st.number_input(f"Impurity Stock Concentration (mg/mL) for {name}", min_value=0.0, value=0.2 if 'IMP' in name else 0.0, key=f"stock_concentration_{name}")
            weight_value = st.number_input(f"Desired Weight (mg) for {name}", min_value=0, value=25 if 'SAM' in name else 20, key=f"weight_{name}")
            v1_value = st.number_input(f"Volume 1 (mL) for {name}", min_value=0, value=50 if 'SAM' in name else 100, key=f"v1_{name}")
            p1_value = st.number_input(f"Pipetted Volume 1 (mL) for {name}", min_value=0, value=5, key=f"p1_{name}")
            v2_value = st.number_input(f"Volume 2 (mL) for {name}", min_value=0, value=10, key=f"v2_{name}")
            p2_value = st.number_input(f"Pipetted Volume 2 (mL) for {name}", min_value=0, value=5, key=f"p2_{name}")
            v3_value = st.number_input(f"Volume 3 (mL) for {name}", min_value=0, value=5, key=f"v3_{name}")
            # Creating a new df with the entered values for this row
            new_row = pd.DataFrame({
                'name': [name],
                'range': [range_value],
                'unit': [unit_value],
                'Potency': [potency_value],
                'CF': [cf_value],
                'Stock_imp_mg_mL': [stock_concentration_value],
                'w': [weight_value],
                'V1': [v1_value],
                'p1': [p1_value],
                'V2': [v2_value],
                'p2': [p2_value],
                'V3': [v3_value]
            })
            # Concatenating the new row to the existing df
            df_info = pd.concat([df_info, new_row], ignore_index=True)
    # Creating a DataFrame with the collected inputs
    #df_info = pd.DataFrame(data_entries)
    # Display the final DataFrame with the entered data
    #st.write("The data you entered:")
    #st.dataframe(df_info)
    def add_tc_and_unit(row):
        '''function to add unit to the target values extracted'''
        name = row['name']
        if name in sam_tc_df['Name'].values:
            sam_tc = sam_tc_df[sam_tc_df['Name'] == name]['sam_tc'].values[0]
            sam_tc_unit = sam_tc_df[sam_tc_df['Name'] == name]['sam_tc_unit'].values[0]
            return pd.Series([sam_tc, sam_tc_unit], index=['tc', 'tc_unit'])
        elif name in TC_std_df['TC_std_name'].values:
            tc_std = TC_std_df[TC_std_df['TC_std_name'] == name]['TC_std'].values[0]
            tc_std_unit = TC_std_df[TC_std_df['TC_std_name'] == name]['TC_std_unit'].values[0]
            return pd.Series([tc_std, tc_std_unit], index=['tc', 'tc_unit'])
        elif name in imp_df['imp_name'].values:
            imp_spec = imp_df[imp_df['imp_name'] == name]['imp_spec'].values[0]
            imp_spec_unit = imp_df[imp_df['imp_name'] == name]['imp_spec_unit'].values[0]
            return pd.Series([imp_spec, imp_spec_unit], index=['tc', 'tc_unit'])
        return pd.Series([np.nan, np.nan], index=['tc', 'tc_unit'])
    df_info[['tc', 'tc_unit']] = df_info.apply(add_tc_and_unit, axis=1)
    def add_sam_tc(row, sam_tc_value):
        if row['Category'] in ['sam', 'std']:
            return row['tc']  
            return sam_tc_value  
    # Check the first tc for sam value 
    first_sam_tc = df_sel[df_sel['Category'] == 'sam'].iloc[0]['tc']
    df_sel['sam_tc'] = df_sel.apply(lambda row: add_sam_tc(row, first_sam_tc), axis=1)
    # Merge df_selt adn df_info
    df_sel = pd.merge(df_sel, df_info[['name', 'w','range','unit', 'Potency', 'CF', 'V1','V2','V3','p1', 'p2']], left_on='Name', right_on='name', how='left')
    df_sel.drop('name', axis=1, inplace=True)
    def to_numeric(value):
        '''Function to garantee that all values are numeric and not strings'''
        try:
            return pd.to_numeric(value, errors='coerce')
        except ValueError:
            return None
    def calculate_df_sel(row):
        '''Function to calculate the dilution factor for impurities'''
        if row['Category'] == 'imp':  
            sam_tc=row['sam_tc']
            tc=row['tc']
            w=row['w']
            df_sel_value=w*100/(tc*sam_tc)
            return df_sel_value
        else:
            return None
    df_sel['DF_sel'] = df_sel.apply(calculate_df_sel, axis=1)
    VF = [1, 2, 5, 10, 20, 25, 50, 100, 200, 250, 500, 1000, 2000]
    VF = [v for v in VF if 10 <= v <= 250]
    pipettes = [0.5, 1, 1.5, 2, 2.5, 3, 4, 5, 6, 7, 8, 9, 10, 12, 14, 15, 20, 25, 50]
    pipettes = [p for p in pipettes if 0.5 <= p <= 10]
    def calculate_dilution_factor(vf_comb, pipette_comb):
        '''Function to achieve dilutions factors'''
        vf_product = 1
        pipette_product = 1
        for vf in vf_comb:
            vf_product *= vf
            
        for pipette in pipette_comb:
            pipette_product *= pipette
        
        dilution_factor = vf_product / pipette_product
        return dilution_factor
    def calculate_uncertainty(vf_comb, pipette_comb):
        '''Function to calculate uncertainty'''
        uncertainty = 0
        for i in range(len(pipette_comb)):
            uncertainty += abs(pipette_comb[i] - vf_comb[i + 1]) / vf_comb[i + 1]
        return uncertainty
    def find_best_dilution_combinations(target_factor):
        '''Function to find the best dilutions combinations based on uncertainty'''
        best_combination = None
        best_uncertainty = float('inf')
        for num_vfs in range(2, len(VF)+1):
            vf_combinations = itertools.combinations(VF, num_vfs)
            for vf_comb in vf_combinations:
                pipette_combinations = itertools.combinations(pipettes, num_vfs-1)
                for pipette_comb in pipette_combinations:
                    valid_combination = True
                    for i in range(len(pipette_comb)):
                        if pipette_comb[i] >= vf_comb[i + 1]:
                            valid_combination = False
                            break
                    if valid_combination:
                        dilution_factor = calculate_dilution_factor(vf_comb, pipette_comb)
                        if abs(dilution_factor - target_factor) / target_factor < 0.01:  # 1% de tolerância
                            uncertainty = calculate_uncertainty(vf_comb, pipette_comb)
                            if uncertainty < best_uncertainty:
                                best_uncertainty = uncertainty
                                best_combination = (vf_comb, pipette_comb)
        return best_combination
    def adjust_vf_combination(row, df_selt):
        target_factor = row['DF_sel']
        best_combination = None
        if row['Category'] in ['std', 'sam']:  # For 'std' e 'sam'
            vf_comb = [row['V1'], row['V2'], row['V3']]  
            pipette_comb = [row['p1'], row['p2']]  
            if row['V3'] == 1:
                vf_comb = [row['V2'], row['V1']]
            elif row['V2'] == 1:
                vf_comb = [row['V1']]
            vf_comb = [v for v in vf_comb if isinstance(v, (int, float))]  
            vf_comb = tuple(sorted(vf_comb, reverse=True))  
            best_combination = (target_factor, vf_comb, tuple(pipette_comb))
        elif row['Category'] == 'imp':
            best_combination = find_best_dilution_combinations(target_factor)
            if best_combination:
                vf_comb, pipette_comb = best_combination
                vf_comb = [v for v in vf_comb if isinstance(v, (int, float))]  
                vf_comb = tuple(sorted(vf_comb, reverse=True))  
                best_combination = (target_factor, vf_comb, tuple(pipette_comb))
            else:
                best_combination = (target_factor, None, None)
        if not best_combination:
            best_combination = (target_factor, (), ())
        return best_combination
    best_combinations = []
    for index, row in df_sel.iterrows():
        best_combination = adjust_vf_combination(row, df_sel)
        best_combinations.append(best_combination)
    df_sel['Best_VF_combination'] = [comb[1] if comb[1] else () for comb in best_combinations]
    df_sel['Best_Pipette_combination'] = [comb[2] if comb[2] else () for comb in best_combinations]

    # Função para realizar as substituições conforme a descrição
    def update_columns_with_combinations(row, df):
        '''Function to update based on category'''
        if row['Category'] in ['imp', 'std', 'sam']:
            vf_combination = row['Best_VF_combination']
            p_combination = row['Best_Pipette_combination']
            row['V1'] = vf_combination[0] if len(vf_combination) > 0 else ''
            row['V2'] = vf_combination[1] if len(vf_combination) > 1 else ''
            row['V3'] = vf_combination[2] if len(vf_combination) > 2 else ''
            row['p1'] = p_combination[0] if len(p_combination) > 0 else ''
            row['p2'] = p_combination[1] if len(p_combination) > 1 else ''
        elif row['Category'] == 'excip':
            app_name = row['app']
            app_row = df[df['Name'] == app_name].iloc[0]
            row['V1'] = app_row['Best_VF_combination'][0] if len(app_row['Best_VF_combination']) > 0 else ''
            row['V2'] = app_row['Best_VF_combination'][1] if len(app_row['Best_VF_combination']) > 1 else ''
            row['V3'] = app_row['Best_VF_combination'][2] if len(app_row['Best_VF_combination']) > 2 else ''
            row['p1'] = app_row['Best_Pipette_combination'][0] if len(app_row['Best_Pipette_combination']) > 0 else ''
            row['p2'] = app_row['Best_Pipette_combination'][1] if len(app_row['Best_Pipette_combination']) > 1 else ''
        return row
    df_sel = df_sel.apply(lambda row: update_columns_with_combinations(row, df_sel), axis=1)    
    def update_w_for_excip(row):
        if row['Category'] == 'excip' and pd.isnull(row['w']):
            V1 = pd.to_numeric(row['V1'], errors='coerce') if pd.notnull(row['V1']) else 1
            V2 = pd.to_numeric(row['V2'], errors='coerce') if pd.notnull(row['V2']) else 1
            V3 = pd.to_numeric(row['V3'], errors='coerce') if pd.notnull(row['V3']) else 1
            p1 = pd.to_numeric(row['p1'], errors='coerce') if pd.notnull(row['p1']) else 1
            p2 = pd.to_numeric(row['p2'], errors='coerce') if pd.notnull(row['p2']) else 1
            V1 = V1 if not pd.isna(V1) else 1
            V2 = V2 if not pd.isna(V2) else 1
            V3 = V3 if not pd.isna(V3) else 1
            p1 = p1 if not pd.isna(p1) else 1
            p2 = p2 if not pd.isna(p2) else 1
            row['w'] = (V1 * V2 * V3) / (p1 * p2) * row['tc']
        return row
    df_selt = df_sel.apply(update_w_for_excip, axis=1)
    df_selt['w_tc'] = df_selt['w'] / (df_selt['Potency'] * df_selt['CF'])
    #st.title('Table for excip')
    #st.write(df_selt)
    def calculate_length_of_tc(row):
        return len(str(row['tc']).replace('.', '')) + (1 if '.' in str(row['tc']) else 0)
    df_selt['tc_length'] = df_selt.apply(calculate_length_of_tc, axis=1)
    def calculate_tc_min_max(row):
        tc_min = round(row['tc'] - (4 / (10 ** (row['tc_length'] - (0 if row['tc_length'] == 1 else 1)))), row['tc_length'] + 1)
        tc_max = round(row['tc'] + (4 / (10 ** (row['tc_length'] - (0 if row['tc_length'] == 1 else 1)))), row['tc_length'] + 1)
        return pd.Series([tc_min, tc_max]) 
    df_selt[['tc_min', 'tc_max']] = df_selt.apply(calculate_tc_min_max, axis=1)
    def count_decimal_places(x):
        if isinstance(x, float):
            return len(str(x).split('.')[1]) if '.' in str(x) else 0
        return 0
    def round_to_tc_decimal_places(row):
        decimal_places = count_decimal_places(row['tc']) 
        w_tc_min = None
        w_tc_max = None
        if row['Category'] in ['std', 'sam']:
            if row['unit'] == 'mg':
                w_tc_min = round(row['w'] - row['range'], 2)
                w_tc_max = round(row['w'] + row['range'], 2)
            elif row['unit'] == '%':
                w_tc_min = round(row['w'] - (row['w_tc'] * row['range'] / 100), 2)
                w_tc_max = round(row['w'] + (row['w_tc'] * row['range'] / 100), 2)
        else:
            w_tc_min = row['tc_min'] * row['w_tc'] / row['tc']
            w_tc_max = row['tc_max'] * row['w_tc'] / row['tc']
            w_tc_min = round(w_tc_min, 2)
            w_tc_max = round(w_tc_max, 2)
        return pd.Series([w_tc_min, w_tc_max])
    #st.write(df_selt)
    df_selt[['w_tc_min', 'w_tc_max']] = df_selt.apply(round_to_tc_decimal_places, axis=1)
    def calculate_D_values(row, df_selt):
        potency = row['Potency'] if pd.notnull(row['Potency']) else 1
        cf = row['CF'] if pd.notnull(row['CF']) else 1
        V1 = pd.to_numeric(row['V1'], errors='coerce') if pd.notnull(row['V1']) else 1
        V2 = pd.to_numeric(row['V2'], errors='coerce') if pd.notnull(row['V2']) else 1
        V3 = pd.to_numeric(row['V3'], errors='coerce') if pd.notnull(row['V3']) else 1
        p1 = pd.to_numeric(row['p1'], errors='coerce') if pd.notnull(row['p1']) else 1
        p2 = pd.to_numeric(row['p2'], errors='coerce') if pd.notnull(row['p2']) else 1
        if row['Category'] in ['imp', 'excip']:
            weight = row['w_tc']  
        else:  
            weight = row['w']
        D1 = (weight*potency * cf) / V1
        if row['Category'] == 'imp':
            sam_tc_value = df_selt[df_selt['Category'] == 'sam']['tc'].iloc[0] if not df_selt[df_selt['Category'] == 'sam'].empty else 1
            D1 = D1 * 100 / sam_tc_value
        D2 = D1 * p1/V2 
        D3 = D2 * p2/V3 
        return pd.Series([D1, D2, D3])
    df_selt[['D1', 'D2', 'D3']] = df_selt.apply(calculate_D_values, axis=1, df_selt=df_selt)
    colunas_necessarias = ["V1", "V2", "V3", "p1", "p2", "p3"]
    for col in colunas_necessarias:
        if col not in df_selt.columns:
            df_selt[col] = 1
    df_selt[colunas_necessarias] = df_selt[colunas_necessarias].replace("", pd.NA)
    df_selt[colunas_necessarias] = df_selt[colunas_necessarias].apply(pd.to_numeric, errors="coerce")
    df_selt = df_selt.fillna(1)
    if (df_selt["p1"] * df_selt["p2"]).eq(0).any():
        raise ValueError("Erro: O denominador contém zero, o que causaria uma divisão por zero.")
    df_selt["DF_final"] = (df_selt["V1"] * df_selt["V2"] * df_selt["V3"]) / (df_selt["p1"] * df_selt["p2"])
    def add_tc_columns_from_spiking(row):
        if isinstance(row['tc_spiking'], list) and row['tc_spiking']:
            for i, tc_value in enumerate(row['tc_spiking']):
                column_name = f"stock_spiking_solution_{i + 1}_tc"
                row[column_name] = tc_value  
        return row
    df_selt2 = df_selt.apply(add_tc_columns_from_spiking, axis=1)
    def round_to_tc_decimal_places(tc, D2):
        decimal_places = len(str(tc).split('.')[1]) if '.' in str(tc) else 0  
        return round(D2, decimal_places)
    def add_stock_spiking_columns(row, df_selt):
        if row['Category'] == 'imp' or (isinstance(row['spiked'], list) and row['spiked']):
            spiking_values = []
            spiking_names = []
            for i, impurity in enumerate(row['spiked']):
                column_name = f"stock_spiking_solution_{i + 1}"
                column_name_with_name = f"stock_spiking_solution_{i + 1}_name"
                column_name_tc = f"stock_spiking_solution_{i + 1}_tc" 
                matching_row = df_selt[df_selt['Name'] == impurity]
                if not matching_row.empty:
                    rounded_D2 = round_to_tc_decimal_places(row[column_name_tc], matching_row['D2'].iloc[0])
                    if rounded_D2 <= 1.5 * row[column_name_tc]:
                        spiking_values.append("D1")
                    else:
                        spiking_values.append("D2")
                    spiking_names.append(matching_row['Name'].iloc[0])
                else:
                    spiking_values.append("D2")
                    spiking_names.append(None)
                row[column_name] = spiking_values[i]
                row[column_name_with_name] = spiking_names[i]
        
        return row
    df_selt2 = df_selt2.apply(add_stock_spiking_columns, axis=1, df_selt=df_selt2)
    def add_concentration_column(row, df_selt):
        max_spiked_columns = len(row['spiked']) if isinstance(row['spiked'], list) else 0
        for i in range(max_spiked_columns):
            column_name = f"stock_spiking_solution_{i + 1}"
            column_name_with_name = f"stock_spiking_solution_{i + 1}_name"
            column_name_conc = f"stock_spiking_solution_{i + 1}_name_conc"
            if column_name in row and row[column_name] is not None:
                impurity_name = row[column_name_with_name]
                if impurity_name:
                    matching_row = df_selt[df_selt['Name'] == impurity_name]
                    if not matching_row.empty:
                        column_for_value = row[column_name] 
                        if column_for_value in matching_row.columns:
                            conc_value = matching_row[column_for_value].iloc[0]
                            row[column_name_conc] = conc_value
                        else:
                            row[column_name_conc] = None
                    else:
                        row[column_name_conc] = None
                else:
                    row[column_name_conc] = None
        
        return row
    df_selt2 = df_selt2.apply(add_concentration_column, axis=1, df_selt=df_selt2)
    def add_volume_columns(row):
        if isinstance(row['tc_spiking'], list) and row['tc_spiking']:
            for i in range(len(row['tc_spiking'])):
                tc_column_name = f"stock_spiking_solution_{i + 1}_tc"
                conc_column_name = f"stock_spiking_solution_{i + 1}_name_conc"
                volume_column_name = f"stock_spiking_solution_{i + 1}_volume"
                if tc_column_name in row and conc_column_name in row:
                    tc_value = row[tc_column_name]
                    conc_value = row[conc_column_name]
                    df_final_value = row['DF_final']
                    row[volume_column_name] = (tc_value * df_final_value) / conc_value
                else:
                    row[volume_column_name] = None
        return row
    df_selt2 = df_selt2.apply(add_volume_columns, axis=1)
    #st.write(df_selt2)
    def generate_paragraph_text(row):
        paragraph = f"{row['Name_sol']}:\n"  
        if row['Category'] in ["BLK", "std", "RES"] or (row['Category'] == "sam" and (not row['spiked'] or row['spiked'] == [""])):
            paragraph += "Prepare as per analytical method."
        elif row['Category'] == 'imp':  
            if pd.notnull(row['D2']):
                paragraph += f"\nStock I for {row['Name']} Solution:\nWeight {round(row['w_tc'], 2)} mg ({row['w_tc_min']} mg to {row['w_tc_max']} mg) into {row['V1']} mL and complete to volume."
            else:
                paragraph += f"\n\nWeight {round(row['w_tc'], 2)} mg ({row['w_tc_min']} mg to {row['w_tc_max']} mg) into {row['V1']} mL and complete to volume."
            if pd.notnull(row['D3']):
                paragraph += f"\n\nStock II for {row['Name']} Solution:\nDilute {row['p1']} mL from Stock I into {row['V2']} mL volumetric flask and complete to volume."
            else:
                paragraph += f"\n\n{row['Name']} at {row['tc']}% solution:\nDilute {row['p1']} mL from Stock I into {row['V2']} mL volumetric flask and complete to volume."
            if pd.notnull(row['D3']):
                paragraph += f"\n\n{row['Name']} at {row['tc']}% solution:\nDilute {row['p2']} mL from Stock II into {row['V3']} mL volumetric flask and complete to volume."
        elif row['Category'] == 'excip':
            if pd.notnull(row['V2']) and row['V2'] == 1 and row['p1'] == 1:
                paragraph += f"Weight {round(row['w'], 2)} ± 10% into {row['V1']} mL volumetric flask and complete to volume."
            else:
                paragraph += f"Stock I for {row['Name']} Solution:\nWeight {round(row['w'], 2)} ± 10% into {row['V1']} mL volumetric flask and complete to volume."
                if pd.notnull(row['V3']) and row['V3'] == 1 and row['p2'] == 1:
                    paragraph += f"\nDilute {row['p2']} mL from Stock I into {row['V2']} mL volumetric flask and complete to volume."
                else:
                    paragraph += f"\nStock II for {row['Name']} Solution:\nDilute {row['p1']} mL from Stock I into {row['V2']} mL volumetric flask and complete to volume."
        elif row['Category'] == 'sam' and row['spiked'] != ['']:
            stock_spiking_phrases = []  
            for i in range(1, 100):  
                col_name = f"stock_spiking_solution_{i}_name"
                col_volume = f"stock_spiking_solution_{i}_volume"
                col_type = f"stock_spiking_solution_{i}"  
                if col_name in row and pd.notnull(row[col_name]):  
                    solution_type = "Stock I" if row[col_type] == "D1" else "Stock II"
                    stock_spiking_phrases.append(f"{row[col_volume]} mL of {solution_type} solution for {row[col_name]}")
            stock_spiking_text = ""
            if stock_spiking_phrases:
                if len(stock_spiking_phrases) == 1:
                    stock_spiking_text = stock_spiking_phrases[0]
                else:
                    stock_spiking_text = ", ".join(stock_spiking_phrases[:-1]) + " and " + stock_spiking_phrases[-1]
            if pd.notnull(row['V2']) and row['V2'] == 1 and row['p1'] == 1:
                paragraph += f"\nWeight {round(row['w'], 2)} ± 10% into {row['V1']} mL volumetric flask"
                if stock_spiking_text:
                    paragraph += f", add {stock_spiking_text}"
                paragraph += " and complete to volume."
            else:  
                paragraph += f"Stock I for {row['Name']} spiked solution:\nWeight {round(row['w'], 2)} ± 10% into {row['V1']} mL volumetric flask"
                if stock_spiking_text:
                    paragraph += f", add {stock_spiking_text}"
                paragraph += " and complete to volume."
                if pd.notnull(row['V3']) and row['V3'] != 1 and row['p2'] != 1:
                    paragraph += f"\nDilute {row['p2']} mL from Stock I into {row['V2']} mL volumetric flask and complete to volume."
                else:
                    paragraph += f"\nStock II for {row['Name']} Solution:\nDilute {row['p1']} mL from Stock I into {row['V2']} mL volumetric flask and complete to volume."
        return paragraph
    paragraphs = df_selt2.apply(generate_paragraph_text, axis=1)
    st.title("Benchwork Plan for Method Validation")
    for paragraph in paragraphs:
        st.text(paragraph + "\n")
else:
    st.warning("Please, upload the document in word format.")
